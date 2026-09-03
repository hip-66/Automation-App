import time
import os
import sys
import datetime
import logging
import ctypes
from docx import Document
from docx.shared import Inches
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from PIL import ImageGrab

# --- Configuration ---
# Never hardcoded: PSAUTO_USERNAME/PASSWORD (explicit override from the app's
# UI) wins; otherwise PSAUTO_DEFAULT_USERNAME/PASSWORD (the app's encrypted
# .env default) is used; a standalone run with neither set prompts instead.
USERNAME = os.environ.get("PSAUTO_USERNAME", "").strip() or os.environ.get("PSAUTO_DEFAULT_USERNAME", "").strip()
PASSWORD = os.environ.get("PSAUTO_PASSWORD", "") or os.environ.get("PSAUTO_DEFAULT_PASSWORD", "")
if not USERNAME or not PASSWORD:
    import getpass
    if not USERNAME:
        USERNAME = input("iDRAC username: ").strip()
    if not PASSWORD:
        PASSWORD = getpass.getpass("iDRAC password: ")

def _resolve_chromedriver():
    """Prefer the exact driver handed over by the PS Automation app (env var).
    Otherwise search upward for the app's Chromedrivers folder and pick the
    newest version in it - so manual runs also get a sane driver."""
    env_path = os.environ.get("CHROMEDRIVER_PATH")
    if env_path and os.path.exists(env_path):
        return env_path
    search_dir = os.path.dirname(os.path.abspath(__file__))
    for _ in range(6):
        drivers_root = os.path.join(search_dir, "Chromedrivers")
        if os.path.isdir(drivers_root):
            import re
            candidates = []
            for root, _dirs, files in os.walk(drivers_root):
                for fname in files:
                    if fname.lower().startswith("chromedriver") and fname.lower().endswith(".exe"):
                        rel = os.path.relpath(os.path.join(root, fname), drivers_root)
                        m = re.search(r"(\d+)", rel)
                        candidates.append((int(m.group(1)) if m else -1, os.path.join(root, fname)))
            if candidates:
                candidates.sort(key=lambda c: c[0], reverse=True)
                return candidates[0][1]
        search_dir = os.path.dirname(search_dir)
    return env_path or r"C:\Scripts\chromedriver.exe"

CHROMEDRIVER_PATH = _resolve_chromedriver()

# Configure Logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger()

def _find_window_containing(substr):
    """Chrome appends ' - Google Chrome' (or similar) to the page title, so
    an exact FindWindowW match on our marker always fails - enumerate
    top-level windows and match by substring instead."""
    user32 = ctypes.windll.user32
    found = []
    EnumWindowsProc = ctypes.WINFUNCTYPE(ctypes.c_bool, ctypes.c_int, ctypes.c_int)

    def _callback(hwnd, _lparam):
        length = user32.GetWindowTextLengthW(hwnd)
        if length:
            buf = ctypes.create_unicode_buffer(length + 1)
            user32.GetWindowTextW(hwnd, buf, length + 1)
            if substr in buf.value:
                found.append(hwnd)
        return True

    user32.EnumWindows(EnumWindowsProc(_callback), 0)
    return found[0] if found else 0

def _restore_title(driver, title):
    """Put a previously-saved document.title back (no-op if none was saved).
    Best-effort: a failure here must never abort the run."""
    if title is None:
        return
    try:
        driver.execute_script("document.title = arguments[0];", title)
    except Exception:
        pass

def bring_chrome_to_foreground(driver):
    """Force the Selenium-controlled Chrome window to the OS foreground.
    Windows blocks background processes from stealing real input focus by
    default (the exact situation here: PS Automation itself runs hidden),
    so this combines two techniques for reliability: (1) SetWindowPos with
    HWND_TOPMOST, which forces the window to the top of the Z-order
    regardless of focus rules - this alone is enough for screenshots, since
    they just capture whatever is visually on top; (2) AttachThreadInput,
    which lets SetForegroundWindow succeed even while another window
    currently holds real/recent user input. Tags the page with a unique
    title first so the right window is found even with other Chrome
    windows open."""
    original_title = None
    try:
        # Remember the page's REAL title before tagging it, so we can put it
        # back the instant the window is located. Both the hostname detection
        # and the browser tab captured in screenshots read document.title -
        # the marker must never linger there (that leak is what produced bogus
        # "PSAUTO_FOCUS_..." server names).
        try:
            original_title = driver.title
        except Exception:
            original_title = None

        marker = f"PSAUTO_FOCUS_{os.getpid()}"
        driver.execute_script(f"document.title = '{marker}';")
        user32 = ctypes.windll.user32
        kernel32 = ctypes.windll.kernel32
        hwnd = 0
        for _ in range(20):
            hwnd = _find_window_containing(marker)
            if hwnd:
                break
            time.sleep(0.15)

        # The marker has served its purpose (or we gave up finding it) -
        # restore the real title NOW, long before hostname detection or any
        # screenshot runs.
        _restore_title(driver, original_title)
        original_title = None  # already restored; don't restore again in finally

        if not hwnd:
            logger.warning("Could not locate the Chrome window to bring it to the foreground.")
            return

        SW_SHOWMAXIMIZED, HWND_TOPMOST, HWND_NOTOPMOST = 3, -1, -2
        SWP_NOMOVE, SWP_NOSIZE, SWP_SHOWWINDOW = 0x0002, 0x0001, 0x0040

        # Show the window MAXIMIZED - NOT SW_RESTORE, which would un-maximize the
        # window that --start-maximized / maximize_window() just maximized,
        # leaving a small windowed Chrome captured in the screenshots.
        user32.ShowWindow(hwnd, SW_SHOWMAXIMIZED)
        user32.SetWindowPos(hwnd, HWND_TOPMOST, 0, 0, 0, 0, SWP_NOMOVE | SWP_NOSIZE | SWP_SHOWWINDOW)

        current_tid = kernel32.GetCurrentThreadId()
        fg_hwnd = user32.GetForegroundWindow()
        fg_tid = user32.GetWindowThreadProcessId(fg_hwnd, None) if fg_hwnd else 0
        target_tid = user32.GetWindowThreadProcessId(hwnd, None)

        attached_fg = attached_target = False
        if fg_tid and fg_tid != current_tid:
            attached_fg = bool(user32.AttachThreadInput(current_tid, fg_tid, True))
        if target_tid and target_tid != current_tid:
            attached_target = bool(user32.AttachThreadInput(current_tid, target_tid, True))
        try:
            user32.BringWindowToTop(hwnd)
            user32.SetForegroundWindow(hwnd)
        finally:
            if attached_fg:
                user32.AttachThreadInput(current_tid, fg_tid, False)
            if attached_target:
                user32.AttachThreadInput(current_tid, target_tid, False)

        user32.SetWindowPos(hwnd, HWND_NOTOPMOST, 0, 0, 0, 0, SWP_NOMOVE | SWP_NOSIZE)
    except Exception as e:
        logger.warning(f"Could not bring Chrome window to foreground: {e}")
    finally:
        # If we bailed out before the mid-function restore (e.g. an exception
        # right after tagging the title), make sure the marker still never
        # survives this call.
        _restore_title(driver, original_title)

def get_screen_resolution():
    """Gets the resolution of the primary monitor only."""
    try:
        user32 = ctypes.windll.user32
        # This line ensures we get the real resolution even if Windows Scaling is > 100%
        user32.SetProcessDPIAware()
        width = user32.GetSystemMetrics(0)
        height = user32.GetSystemMetrics(1)
        return width, height
    except Exception as e:
        logger.error(f"Error getting screen resolution: {e}")
        return 1920, 1080 # Fallback default

def take_full_screen_shot(filename):
    """
    Captures ONLY the primary monitor (including taskbar/clock).
    Ignores the second monitor using bbox.
    """
    try:
        # Give a split second for rendering to finish
        time.sleep(1)
        
        # 1. Get dimensions of the primary monitor
        width, height = get_screen_resolution()
        
        # 2. Define the bounding box (0,0 is top-left of primary screen)
        # This prevents capturing the extended second monitor
        screenshot = ImageGrab.grab(bbox=(0, 0, width, height))
        
        screenshot.save(filename)
        return True
    except Exception as e:
        logger.error(f"Error taking full screenshot: {e}")
        return False

def add_to_word(doc, title, image_path):
    """Adds a title, timestamp, and image to the Word document."""
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_heading(title, level=2)
    doc.add_paragraph(f"Taken at: {now}") 
    try:
        doc.add_picture(image_path, width=Inches(6.5))
    except Exception as e:
        doc.add_paragraph(f"Error adding image: {e}")
    doc.add_paragraph("\n")

def force_click_element(driver, element):
    """Executes a JavaScript click to bypass overlays or visibility issues."""
    driver.execute_script("arguments[0].click();", element)

def wait_for_text_in_body(driver, text, timeout=15):
    """Waits until specific text appears in the body tag to confirm page load."""
    try:
        WebDriverWait(driver, timeout).until(
            EC.text_to_be_present_in_element((By.TAG_NAME, "body"), text)
        )
        return True
    except:
        return False

def process_single_server(server, doc):
    """
    Connects to a single server, extracts its name, takes screenshots,
    and returns the status and the extracted name.
    """
    initial_name = server['name']
    url = server['url']
    
    # Placeholder name in case extraction fails
    current_server_name = initial_name 
    
    logger.info(f"==========================================")
    logger.info(f"Processing URL: {url}")

    options = webdriver.ChromeOptions()
    options.add_argument("--start-maximized")
    options.add_argument("--ignore-certificate-errors")
    options.add_argument("--allow-running-insecure-content")
    options.add_argument("--no-proxy-server")
    # Remove the "Chrome is being controlled by automated test software" info
    # bar (and the automation extension/flag that produce it) so it never
    # appears in the screenshots captured for the report.
    options.add_experimental_option("excludeSwitches", ["enable-automation"])
    options.add_experimental_option("useAutomationExtension", False)
    options.add_argument("--disable-blink-features=AutomationControlled")

    service = Service(executable_path=CHROMEDRIVER_PATH)
    driver = webdriver.Chrome(service=service, options=options)

    # Force the window to fill the screen. This is the reliable, canonical way -
    # the --start-maximized flag alone is not always honored, and the screenshots
    # grab the whole primary monitor, so a maximized window is what makes them
    # look right (rather than a small window with desktop around it).
    try:
        driver.maximize_window()
    except Exception:
        pass

    # Bring window to front
    driver.switch_to.window(driver.current_window_handle)
    bring_chrome_to_foreground(driver)

    try:
        driver.set_page_load_timeout(60)
        driver.get(url)
        bring_chrome_to_foreground(driver)  # retry once a real page is loaded

        # --- DYNAMIC HOSTNAME EXTRACTION ---
        try:
            logger.info("Attempting to extract hostname from page title...")
            # Wait for a REAL (non-marker) title. The foreground helper briefly
            # tags the tab as PSAUTO_FOCUS_<pid>; guard against ever reading
            # that leftover marker as if it were the server's hostname.
            WebDriverWait(driver, 10).until(
                lambda d: len(d.title) > 1 and not d.title.startswith("PSAUTO_FOCUS")
            )
            page_title = driver.title

            if " - " in page_title:
                extracted_name = page_title.split(" - ")[0].strip()
            else:
                extracted_name = page_title.strip()

            if extracted_name and not extracted_name.startswith("PSAUTO_FOCUS"):
                current_server_name = extracted_name
                logger.info(f"Hostname DETECTED: {current_server_name}")
            else:
                logger.warning("Page title did not resolve to a real hostname; using IP/Default name.")

            server['extracted_name'] = current_server_name

        except Exception as title_error:
            logger.warning(f"Could not extract dynamic hostname: {title_error}. Using IP/Default.")

        # Add Heading to Doc using the NEW detected name
        doc.add_heading(f"Server: {current_server_name} - {url}", level=1)
        
        # --- LOGIN ---
        logger.info(f"Attempting Login for {current_server_name}...")
        try:
            wait = WebDriverWait(driver, 20)
            user_field = wait.until(EC.element_to_be_clickable((By.NAME, "username")))
            user_field.clear()
            user_field.send_keys(USERNAME)
            time.sleep(0.5)

            pass_field = driver.find_element(By.NAME, "password")
            driver.execute_script("arguments[0].value = arguments[1];", pass_field, PASSWORD)
            pass_field.send_keys(" ") 
            pass_field.send_keys(Keys.BACKSPACE)
            time.sleep(0.5)
            pass_field.send_keys(Keys.RETURN)

            logger.info("Login submitted. Waiting 15s...")
            time.sleep(15)
            
            # Close Popups
            try:
                popups = driver.find_elements(By.XPATH, "//*[contains(text(), 'Later') or contains(text(), 'Keep Default')]")
                for p in popups:
                    if p.is_displayed():
                        p.click()
            except:
                pass

        except Exception as e:
            logger.error(f"Login process error: {e}")

        # Safe filename for temp images
        safe_name = "".join(x for x in current_server_name if x.isalnum() or x in "-_")

        # =========================================================
        # STEP 1: DASHBOARD
        # =========================================================
        logger.info("Taking Dashboard Screenshot...")
        img_name = f"temp_{safe_name}_dashboard.png"
        
        if take_full_screen_shot(img_name):
            add_to_word(doc, "Dashboard", img_name)
            os.remove(img_name)

        # =========================================================
        # STEP 2: STORAGE (Navigation)
        # =========================================================
        logger.info("Navigating to Storage -> Overview...")
        try:
            storage_link = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.XPATH, "//ul[contains(@class,'navbar')]//a[contains(., 'Storage')]"))
            )
            storage_link.click()
            time.sleep(2)
            try:
                overview_link = driver.find_element(By.XPATH, "//a[contains(text(), 'Overview') and @class='visible']")
                overview_link.click()
            except:
                pass 

            if wait_for_text_in_body(driver, "Storage"):
                logger.info("Storage page loaded.")
            else:
                logger.warning("Storage page might not have loaded correctly.")

            time.sleep(3) 

            # =========================================================
            # STEP 3: STORAGE TABS
            # =========================================================
            tabs_config = [
                {"name": "Summary", "xpath": "//div[contains(@id, 'summary')]", "verify": "Summary"},
                {"name": "Controllers", "xpath": "//div[contains(@id, 'controllers')]", "verify": "Controller"},
                {"name": "Physical Disks", "xpath": "//div[contains(@id, 'pdisks')]", "verify": "Physical Disks"},
                {"name": "Virtual Disks", "xpath": "//div[contains(@id, 'vdisks')]", "verify": "Virtual Disks"}
            ]

            for tab in tabs_config:
                tab_name = tab["name"]
                verify_text = tab["verify"]
                tab_xpath = tab["xpath"]
                
                logger.info(f"--- Processing Tab: {tab_name} ---")
                try:
                    wait = WebDriverWait(driver, 5)
                    tab_elem = wait.until(EC.presence_of_element_located((By.XPATH, tab_xpath)))
                    force_click_element(driver, tab_elem)
                    logger.info(f"Clicked {tab_name} (JS Force).")
                    
                    time.sleep(4) 

                    if wait_for_text_in_body(driver, verify_text, timeout=10):
                        logger.info(f"VERIFIED: Page updated to {tab_name}.")
                        time.sleep(1) 
                        
                        # Dynamic multi-page handler - both Physical Disks AND
                        # Virtual Disks can span multiple pages on a server with
                        # enough disks/virtual drives; previously only Physical
                        # Disks paginated and Virtual Disks always grabbed just
                        # the first page.
                        if tab_name in ("Physical Disks", "Virtual Disks"):
                            logger.info("Detecting total pages available...")

                            # Find all potential numeric page numbers inside the active page view area
                            page_numbers = set()
                            for el in driver.find_elements(By.TAG_NAME, "a"):
                                try:
                                    text = el.text.strip()
                                    if text.isdigit():
                                        page_numbers.add(int(text))
                                except:
                                    continue

                            total_pages = max(page_numbers) if page_numbers else 1
                            logger.info(f"Discovered total pages count: {total_pages}")

                            # Iteratively process each page index independently to prevent stale DOM objects
                            for p_idx in range(1, total_pages + 1):
                                logger.info(f"Processing {tab_name} - Page {p_idx}/{total_pages}")
                                
                                if p_idx > 1:
                                    try:
                                        # Locate the anchor element using link text dynamically at runtime
                                        page_btn = driver.find_element(By.LINK_TEXT, str(p_idx))
                                        force_click_element(driver, page_btn)
                                        time.sleep(4) # Allow data rows UI to update completely
                                    except Exception as click_err:
                                        logger.error(f"Failed navigating to page {p_idx} natively. Trying general XPATH lookups: {click_err}")
                                        try:
                                            # Fallback XPath approach if exact link text matching hits anomalies
                                            page_btn = driver.find_element(By.XPATH, f"//a[text()='{p_idx}' or normalize-space(text())='{p_idx}']")
                                            force_click_element(driver, page_btn)
                                            time.sleep(4)
                                        except Exception as fallback_err:
                                            logger.critical(f"Aborting pagination click loop for page {p_idx}: {fallback_err}")
                                            continue
                                
                                # Process full screen snapshot output per active table page configuration
                                img_name = f"temp_{safe_name}_{tab_name}_Page_{p_idx}.png"
                                if take_full_screen_shot(img_name):
                                    add_to_word(doc, f"Storage - {tab_name} (Page {p_idx})", img_name)
                                    os.remove(img_name)
                        else:
                            # Standard rendering pattern execution for single-grid target tabs
                            img_name = f"temp_{safe_name}_{tab_name}.png"
                            if take_full_screen_shot(img_name):
                                add_to_word(doc, f"Storage - {tab_name}", img_name)
                                remove_target_temp_asset = True
                                os.remove(img_name)
                    else:
                        logger.error(f"VALIDATION FAILED: {tab_name} text missing.")
                        img_name = f"temp_{safe_name}_{tab_name}_fallback.png"
                        if take_full_screen_shot(img_name):
                            add_to_word(doc, f"Storage - {tab_name} (Fallback)", img_name)
                            os.remove(img_name)

                except Exception as e:
                    logger.error(f"Failed to process tab {tab_name}: {e}")

        except Exception as e:
            logger.error(f"Critical error in Storage section: {e}")

        # =========================================================
        # STEP 4: iDRAC SETTINGS
        # =========================================================
        logger.info("Navigating to iDRAC Settings...")
        try:
            settings_xpath = "//*[@id='settings']"
            elem = WebDriverWait(driver, 10).until(EC.element_to_be_clickable((By.XPATH, settings_xpath)))
            force_click_element(driver, elem)
            logger.info("Clicked iDRAC Settings link.")
            time.sleep(5) 
            
            if wait_for_text_in_body(driver, "Connectivity") or wait_for_text_in_body(driver, "Services") or wait_for_text_in_body(driver, "Settings"):
                img_name = f"temp_{safe_name}_settings.png"
                if take_full_screen_shot(img_name):
                    add_to_word(doc, "iDRAC Settings", img_name)
                    os.remove(img_name)
            else:
                logger.error("FAILED: iDRAC Settings page did not load correctly.")

        except Exception as e:
            logger.error(f"Failed iDRAC Settings navigation: {e}")

        return "Success", "Completed successfully", current_server_name

    except Exception as e:
        logger.error(f"CRITICAL ERROR on {current_server_name}: {e}")
        doc.add_paragraph(f"Connection Failed: {e}")
        return "Failed", str(e), current_server_name
    
    finally:
        driver.quit()
        logger.info(f"Finished {current_server_name}.")

def main():
    if not os.path.exists(CHROMEDRIVER_PATH):
        logger.error(f"Error: chromedriver.exe not found at {CHROMEDRIVER_PATH}")
        return

    while True:
        print("\n" + "="*50)
        print("    iDRAC Auto-Screenshot Tool Configuration")
        print("="*50)
        print("Select Mode:")
        print("1. Sequence Mode (Base IP + Start Suffix + Count)")
        print("2. Specific IP List Mode (Enter full IPs manually)")
        print("3. Exit")
        print("="*50)
        
        mode_choice = input("Enter choice (1, 2 or 3): ").strip()
        servers_list = []
        
        if mode_choice == "3":
            print("Exiting program.")
            return

        elif mode_choice == "1":
            print("\n--- Mode 1 Selected ---")
            base_ip_prefix = input("Enter Base IP (e.g., 192.168.0): ").strip()
            
            if base_ip_prefix.count('.') < 1:
                print("\n[!] ERROR: Invalid Base IP format. It must look like '192.168.0'.")
                input("Press Enter to return to menu...")
                continue 

            try:
                start_suffix = int(input("Enter Starting Suffix (e.g., 120): "))
                server_count = int(input("Enter number of servers to scan: "))
            except ValueError:
                print("\n[!] ERROR: Suffix and Count must be valid integer numbers.")
                input("Press Enter to return to menu...")
                continue 

            for i in range(server_count):
                current_suffix = start_suffix + i
                ip = f"{base_ip_prefix}.{current_suffix}"
                server_obj = {
                    "name": f"Server_{ip}", 
                    "url": f"https://{ip}/"
                }
                servers_list.append(server_obj)
            break 

        elif mode_choice == "2":
            print("\n--- Mode 2 Selected ---")
            print("Enter full IP addresses (e.g., 10.201.236.167).")
            print("Type 'done' or press Enter on an empty line to finish.")
            
            counter = 1
            while True:
                user_ip = input(f"Enter IP #{counter}: ").strip()
                if user_ip.lower() == "done" or user_ip == "":
                    break
                
                if len(user_ip) < 7: 
                    print("Invalid IP length, ignored.")
                    continue

                server_obj = {
                    "name": f"Server_{user_ip}", 
                    "url": f"https://{user_ip}/"
                }
                servers_list.append(server_obj)
                counter += 1
            
            if not servers_list:
                print("\n[!] Warning: No valid IPs entered.")
                input("Press Enter to return to menu...")
                continue 
            break
        else:
            print("\n[!] Invalid choice. Please select 1, 2, or 3.")
            time.sleep(1)
            continue

    print("-" * 50)
    print(f"Starting process for {len(servers_list)} servers...")
    print("-" * 50)

    doc = Document()
    doc.add_heading('iDRAC System Report', 0)

    execution_summary = []
    detected_hostnames = []

    # The output filename is decided up front (safe fallback based on the
    # first target IP) and the SAME file is re-saved after every single
    # server - so if the run is killed/crashes partway through, whatever
    # was already processed is never lost. It gets renamed to a nicer
    # hostname-based name at the very end, once that's known.
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    file_prefix = "iDRAC_Report"
    first_ip = servers_list[0]['url'].replace("https://", "").replace("/", "")
    dynamic_output_file = f"{file_prefix}_{first_ip}_{timestamp}.docx"

    def safe_save():
        try:
            doc.save(dynamic_output_file)
            logger.info(f"Progress saved incrementally to: {dynamic_output_file}")
        except PermissionError:
            logger.error(f"ERROR: Could not save '{dynamic_output_file}'. Is the file open?")

    for server in servers_list:
        status, msg, final_name = process_single_server(server, doc)

        if status == "Success" and final_name != server['name']:
            detected_hostnames.append(final_name)

        execution_summary.append({
            "name": final_name,
            "url": server['url'],
            "status": status,
            "msg": msg
        })

        safe_save()

    if detected_hostnames:
        first_name = detected_hostnames[0]
        system_name = first_name.split("-")[0] if "-" in first_name else first_name
        renamed_file = f"{system_name}_{file_prefix}_{timestamp}.docx"
        try:
            if os.path.exists(dynamic_output_file):
                os.replace(dynamic_output_file, renamed_file)
                dynamic_output_file = renamed_file
        except Exception as e:
            logger.warning(f"Could not rename report to hostname-based filename: {e}")

    logger.info(f"Report saved: {dynamic_output_file}")

    print("\n\n")
    print("="*80)
    print(f"{'FINAL EXECUTION SUMMARY':^80}")
    print("="*80)
    print(f"{'Server Name':<25} | {'IP Address':<25} | {'Status':<10} | {'Note'}")
    print("-" * 80)
    
    for item in execution_summary:
        clean_ip = item['url'].replace("https://", "").replace("/", "")
        note = (item['msg'][:20] + '..') if len(item['msg']) > 20 else item['msg']
        print(f"{item['name']:<25} | {clean_ip:<25} | {item['status']:<10} | {note}")
    
    print("="*80)
    print("\n")
    # Pause only when run interactively from a console; when launched by the
    # PS Automation app stdin is a closed pipe and input() would crash with
    # EOFError, wrongly marking a successful run as failed.
    if sys.stdin is not None and sys.stdin.isatty():
        input("Press Enter to exit...")

if __name__ == "__main__":
    main()