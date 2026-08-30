import os
import glob
import pandas as pd
import ssl
import atexit
from datetime import datetime
from pyVmomi import vim
from pyVim.connect import SmartConnect, Disconnect
from colorama import Fore, init
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Initialize Colorama
init(autoreset=True)
validation_results = []

def log_result(category, item, status, details, ip_address="", host_info="", cpu="", ram="", os_name=""):
    """
    Logs result to console and appends to the global validation list for the report.
    """
    # Console Output logic
    color = Fore.GREEN if status == "OK" else (Fore.YELLOW if status == "WARNING" else Fore.RED)
    # Print simplified log to console to keep it clean
    print(f"{color}[{status}] [{category}] {item} | {details[:100]}...")
    
    # Store data for Word Report
    validation_results.append({
        "Category": category, "Item": item, "Status": status,
        "Details": details, "IP": ip_address, "HostInfo": host_info,
        "CPU": cpu, "RAM": ram, "OS": os_name
    })

def get_obj(content, vimtype, name):
    """
    Helper function to find a vSphere object by name.
    """
    obj = None
    container = content.viewManager.CreateContainerView(content.rootFolder, vimtype, True)
    for c in container.view:
        if c.name == name:
            obj = c
            break
    container.Destroy()
    return obj

def set_cell_background(cell, color_hex):
    """
    Sets the background color of a Word table cell.
    """
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

# --- 1. Locate and Read Configuration File ---
def find_excel_file():
    print("--- Select Configuration File ---")
    print("1. Search on Desktop\n2. Enter manual path")
    choice = input("Choose option (1/2): ")
    search_pattern = "*VIM*.xlsm"
    
    if choice == '1':
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        path = os.path.join(desktop, search_pattern)
        files = glob.glob(path) + glob.glob(path.replace(".xlsm", ".xlsx"))
    else:
        folder = input("Enter folder path: ").strip('"')
        path = os.path.join(folder, search_pattern) if os.path.isdir(folder) else folder
        files = glob.glob(path) + glob.glob(path.replace(".xlsm", ".xlsx")) if os.path.isdir(folder) else [path]

    if not files or not os.path.isfile(files[0]):
        print(f"{Fore.RED}No file found.")
        return None
    print(f"{Fore.YELLOW}File found: {files[0]}")
    return files[0]

def read_config(file_path):
    print(f"{Fore.YELLOW}Reading Excel data...")
    try:
        xls = pd.ExcelFile(file_path)
        # Read Sheets
        env_df = pd.read_excel(xls, sheet_name='Environment')
        env_dict = dict(zip(env_df['Parameter Name'].astype(str).str.strip(), env_df['Parameter Value']))
        hosts_df = pd.read_excel(xls, sheet_name='Hosts')
        vms_df = pd.read_excel(xls, sheet_name='VMs')
        return env_dict, hosts_df, vms_df
    except Exception as e:
        print(f"{Fore.RED}Error reading Excel: {e}")
        return None, None, None

# --- 2. Connect to vCenter ---
def connect_to_vcenter(env_config, vms_df):
    host = None
    print(f"{Fore.CYAN}Searching for vCenter IP in 'VMs' sheet...")
    # Find vCenter IP dynamically from the VM list
    vcenter_row = vms_df[vms_df['Computer Name'].astype(str).str.contains('vcenter', case=False, na=False)]
    
    if not vcenter_row.empty:
        if 'Guest IP' in vcenter_row.columns: host = str(vcenter_row.iloc[0]['Guest IP']).strip()
        elif 'IP Address' in vcenter_row.columns: host = str(vcenter_row.iloc[0]['IP Address']).strip()

    # Fallback to Environment sheet
    if not host or host.lower() == 'nan':
        host = env_config.get('vCenter IP') or env_config.get('vCenter Address')
    
    # Last resort: Manual input
    if not host:
        host = input("Enter vCenter IP manually: ").strip()

    print(f"{Fore.YELLOW}Connecting to {host}...")
    # Bypass SSL verification for self-signed certs
    context = ssl._create_unverified_context()
    try:
        si = SmartConnect(host=host, user=env_config.get('vCenter Admin User'), 
                          pwd=env_config.get('vCenter Admin Password'), sslContext=context)
        atexit.register(Disconnect, si)
        print(f"{Fore.GREEN}Connected!")
        return si, host
    except Exception as e:
        print(f"{Fore.RED}Connection failed: {e}")
        return None, host

# --- 3. Validation Logic ---
def validate_environment(si, env_config, hosts_df, vms_df):
    content = si.RetrieveContent()
    
    # --- Environment Checks ---
    print("\n--- Environment Checks ---")
    dc_name = env_config.get('Datacenter Name')
    datacenter = get_obj(content, [vim.Datacenter], dc_name)
    if datacenter: 
        log_result("Environment", "Datacenter", "OK", f"Matched: {dc_name}")
    else: 
        # Fallback to show available DCs
        found_dcs = [c.name for c in content.viewManager.CreateContainerView(content.rootFolder, [vim.Datacenter], True).view]
        log_result("Environment", "Datacenter", "FAIL", f"Missing! (Current: {found_dcs} vs Expected: '{dc_name}')")
        datacenter = content.rootFolder

    cluster_name = env_config.get('Cluster Name')
    cluster = get_obj(content, [vim.ClusterComputeResource], cluster_name)
    if cluster: 
        log_result("Environment", "Cluster", "OK", f"Matched: {cluster_name}")
    else: 
        log_result("Environment", "Cluster", "FAIL", f"Missing! (Expected: '{cluster_name}' not found)")

    # --- Hosts Checks ---
    print("\n--- Hosts Checks ---")
    actual_hosts = {}
    if cluster:
        for h in cluster.host:
            mgmt_ip = "Unknown"
            if h.config.network.vnic:
                for v in h.config.network.vnic:
                    if v.spec.ip.ipAddress: mgmt_ip = v.spec.ip.ipAddress; break
            actual_hosts[mgmt_ip] = h

    for index, row in hosts_df.iterrows():
        excel_ip = str(row['Host IP']).strip()
        excel_name = str(row.get('Host Name', excel_ip)).strip()
        is_enabled = str(row.get('Enabled', 'Yes')).strip().lower() == 'yes'

        if excel_ip in actual_hosts:
            if is_enabled: log_result("Hosts", excel_name, "OK", "Host Online", excel_ip)
            else: log_result("Hosts", excel_name, "WARNING", "Host Found (But Marked Disabled in Excel)", excel_ip)
        else:
            if is_enabled: log_result("Hosts", excel_name, "FAIL", f"Host Missing/Offline! (Expected IP: {excel_ip})", excel_ip)
            else: log_result("Hosts", excel_name, "OK", "Host Offline (As expected)", excel_ip)

    # --- VM Deep Validation ---
    print("\n--- VM Deep Validation ---")
    container = content.viewManager.CreateContainerView(datacenter, [vim.VirtualMachine], True)
    vms_dict = {vm.name: vm for vm in container.view}
    container.Destroy()

    for index, row in vms_df.iterrows():
        vm_name = str(row['Computer Name']).strip()
        is_enabled = str(row.get('Enabled', 'Yes')).strip().lower() == 'yes'
        
        # Get Excel Data
        excel_ip = str(row.get('Guest IP', row.get('IP Address', ''))).strip()
        if excel_ip == 'nan': excel_ip = ""
        excel_host_ip = str(row.get('Host IP', '')).strip()
        if excel_host_ip == 'nan': excel_host_ip = ""
        excel_cpu = row.get('Cores')
        excel_mem_gb = row.get('Memory')
        excel_ds = str(row.get('Datastore', '')).strip()
        
        # Get OS Profile from Excel
        excel_os_profile = str(row.get('OS Profile', '')).strip()
        if excel_os_profile == 'nan': excel_os_profile = ""

        vm = vms_dict.get(vm_name)
        
        # --- Case 1: VM Not Found ---
        if not vm:
            status = "FAIL" if is_enabled else "OK"
            msg = "Defined in Excel, but NOT found in vCenter Inventory" if is_enabled else "Not in vCenter (Matches Excel Disabled)"
            log_result("VMs", vm_name, status, msg, excel_ip)
            continue

        # --- Case 2: VM Found (Validate all parameters) ---
        details_list = []
        critical_fail = False
        
        # Fetch Real Data from vCenter
        real_cpu = vm.config.hardware.numCPU
        real_mem_mb = vm.config.hardware.memoryMB
        real_mem_gb = round(real_mem_mb / 1024, 1) 
        real_os = vm.summary.config.guestFullName if vm.summary.config.guestFullName else "Unknown"

        # 1. Power State Check
        pwr_state = vm.runtime.powerState
        if is_enabled and pwr_state != 'poweredOn':
            details_list.append("FAIL: Power is OFF (Expected: ON)")
            critical_fail = True
        elif not is_enabled and pwr_state == 'poweredOn':
            details_list.append("WARN: Power is ON (Expected: OFF)")

        # 2. IP Address Check
        real_ip = vm.guest.ipAddress if vm.guest.ipAddress else "Unknown"
        if excel_ip and excel_ip != real_ip:
            details_list.append(f"FAIL: IP Mismatch (Current: {real_ip} vs Expected: {excel_ip})")
            critical_fail = True
        
        # 3. Host Affinity Check
        host_msg = ""
        if excel_host_ip:
            runtime_host_ip = "Unknown"
            try:
                h = vm.runtime.host
                if h.config.network.vnic:
                     for v in h.config.network.vnic:
                         if v.spec.ip.ipAddress: 
                             runtime_host_ip = v.spec.ip.ipAddress; break
            except: pass
            
            if runtime_host_ip == excel_host_ip: 
                host_msg = f"Host OK ({runtime_host_ip})"
            else: 
                host_msg = f"Host Mismatch! (Current: {runtime_host_ip} vs Expected: {excel_host_ip})"
                critical_fail = True

        # 4. CPU Check
        if pd.notna(excel_cpu):
            if real_cpu != int(excel_cpu): 
                details_list.append(f"FAIL: CPU Mismatch (Current: {real_cpu} vs Expected: {int(excel_cpu)})")
                critical_fail = True
        
        # 5. RAM Check
        if pd.notna(excel_mem_gb):
            excel_mem_mb = int(excel_mem_gb * 1024)
            # Allow variance of ~4MB
            if abs(real_mem_mb - excel_mem_mb) > 4: 
                details_list.append(f"FAIL: RAM Mismatch (Current: {real_mem_gb}GB vs Expected: {excel_mem_gb}GB)")
                critical_fail = True

        # 6. Datastore Check
        if excel_ds and excel_ds != 'nan':
            vm_datastores = [ds.name for ds in vm.datastore]
            if excel_ds not in vm_datastores:
                details_list.append(f"FAIL: Datastore Mismatch (Current: {vm_datastores} vs Expected: '{excel_ds}')")
                critical_fail = True

        # 7. OS Check (Using 'OS Profile' column)
        if excel_os_profile:
            # Logic: Remove "Profile" from Excel string and check if it exists in vCenter OS string
            # Example: "WindowsProfile" -> "Windows"
            search_term = excel_os_profile.replace("Profile", "").lower()
            
            match_os = False
            if search_term in real_os.lower():
                match_os = True
            # Special case for Linux variants
            elif search_term == "linux" and any(x in real_os.lower() for x in ['red hat', 'centos', 'ubuntu', 'debian']):
                match_os = True
                
            if not match_os:
                details_list.append(f"FAIL: OS Mismatch (Current: '{real_os}' vs Expected Profile: '{excel_os_profile}')")
                critical_fail = True

        # Determine Final Status
        final_status = "OK"
        if critical_fail: final_status = "FAIL"
        elif "WARN" in str(details_list): final_status = "WARNING"
        elif not is_enabled: final_status = "WARNING"

        if not details_list: details_list.append("All Configs Match")
        
        # Join details with newlines for the report
        full_details = "\n".join(details_list)
        
        log_result("VMs", vm_name, final_status, full_details, real_ip, host_msg, 
                   str(real_cpu), f"{real_mem_gb} GB", real_os)

# --- 4. Generate Word Report ---
def generate_word_report(vcenter_address):
    print(f"\n{Fore.YELLOW}Generating Word Report...")
    doc = Document()
    
    # Set Landscape orientation (approximate A4)
    section = doc.sections[0]
    section.page_width = Pt(842) 
    section.page_height = Pt(595)

    doc.add_heading('vSphere Audit Report', 0)
    doc.add_paragraph(f"vCenter: {vcenter_address}\nDate: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # Create Table
    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    headers = ['Category', 'Item', 'IP', 'Host Info', 'CPU', 'RAM', 'OS', 'Status', 'Details']
    
    # Header Styling
    for i, h in enumerate(headers):
        run = table.rows[0].cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    # Populate Data
    for entry in validation_results:
        row = table.add_row().cells
        row[0].text = str(entry['Category'])
        row[1].text = str(entry['Item'])
        row[2].text = str(entry['IP'])
        row[3].text = str(entry['HostInfo'])
        row[4].text = str(entry.get('CPU', ''))
        row[5].text = str(entry.get('RAM', ''))
        row[6].text = str(entry.get('OS', ''))
        row[8].text = str(entry['Details'])

        # Status Cell Styling (Icons & Colors)
        status_cell = row[7]
        p = status_cell.paragraphs[0]
        run = p.add_run()
        
        if entry['Status'] == 'OK':
            run.text = "✔ OK"
            run.font.color.rgb = RGBColor(0, 153, 0) # Green
            run.bold = True
        elif entry['Status'] == 'FAIL':
            run.text = "✖ FAIL"
            run.font.color.rgb = RGBColor(255, 0, 0) # Red
            run.bold = True
            # Highlight entire row Red
            for cell in row: set_cell_background(cell, "FFCCCC")
        elif entry['Status'] == 'WARNING':
            run.text = "⚠ WARNING"
            run.font.color.rgb = RGBColor(255, 140, 0) # Orange
            run.bold = True
            # Highlight entire row Yellow
            for cell in row: set_cell_background(cell, "FFF4CC")
        
        # Adjust Font Size for Content
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    # Save File
    filename = f"Audit_Report_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx"
    try:
        doc.save(filename)
        print(f"{Fore.GREEN}Saved: {filename}")
        os.startfile(filename)
    except Exception as e:
        print(f"{Fore.RED}Error saving report: {e}")

# --- Main Execution ---
def main():
    excel = find_excel_file()
    if not excel: return
    env, hosts, vms = read_config(excel)
    if not env: return
    si, host = connect_to_vcenter(env, vms)
    if si:
        validate_environment(si, env, hosts, vms)
        generate_word_report(host)

if __name__ == "__main__":
    main()