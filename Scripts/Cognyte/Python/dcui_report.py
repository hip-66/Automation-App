import subprocess
import sys
import os
import time
import datetime
import logging
import ctypes  # Library to interact with Windows User32 API

# --- מנגנון בדיקה והתקנה אוטומטית של ספריות חסרות ---
required_libraries = {
    "docx": "python-docx",
    "selenium": "selenium",
    "PIL": "pillow"
}

for lib_import_name, pip_package_name in required_libraries.items():
    try:
        __import__(lib_import_name)
    except ImportError:
        print(f"Library '{pip_package_name}' is missing. Installing it now...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", pip_package_name])
            print(f"Successfully installed '{pip_package_name}'.")
        except Exception as e:
            print(f"Failed to install '{pip_package_name}' automatically. Error: {e}")
            sys.exit(1)

# טעינת הספריות לאחר וידוא התקנה
from docx import Document
from docx.shared import Inches
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from PIL import ImageGrab

# --- Configuration ---
# Never hardcoded: PSAUTO_USERNAME/PASSWORD (explicit override from the app's
# UI) wins; otherwise PSAUTO_DEFAULT_USERNAME/PASSWORD (the app's encrypted
# .env default) is used; a standalone run with neither set prompts instead.
USERNAME = os.environ.get("PSAUTO_USERNAME", "").strip() or os.environ.get("PSAUTO_DEFAULT_USERNAME", "").strip()
PASSWORD = os.environ.get("PSAUTO_PASSWORD", "") or os.environ.get("PSAUTO_DEFAULT_PASSWORD", "")
if not USERNAME or not PASSWORD:
    if sys.stdin is not None and sys.stdin.isatty():
        import getpass
        if not USERNAME:
            USERNAME = input("iDRAC username: ").strip()
        if not PASSWORD:
            PASSWORD = getpass.getpass("iDRAC password: ")
    else:
        print("ERROR: No iDRAC username/password available and there is no console to prompt on.", file=sys.stderr, flush=True)
        sys.exit(1)

def _resolve_chromedriver():
    """Prefer the exact driver handed over by the PS Automation app (env var).
    Otherwise search upward for the app's Chromedrivers folder and pick the
    newest version in it - so manual runs also get a sane driver."""
    env_path = os.environ.get("CHROMEDRIVER_PATH")
    if env_path and os.path.exists(env_path):
        return env_path
    search_dir = os.path.dirname(os.path.abspath(__file__))
    for _ in range(6):
        drivers_root = os.path.join(search_dir, "Chromedrivers")
        if os.path.isdir(drivers_root):
            import re
            candidates = []
            for root, _dirs, files in os.walk(drivers_root):
                for fname in files:
                    if fname.lower().startswith("chromedriver") and fname.lower().endswith(".exe"):
                        rel = os.path.relpath(os.path.join(root, fname), drivers_root)
                        m = re.search(r"(\d+)", rel)
                        candidates.append((int(m.group(1)) if m else -1, os.path.join(root, fname)))
            if candidates:
                candidates.sort(key=lambda c: c[0], reverse=True)
                return candidates[0][1]
        search_dir = os.path.dirname(search_dir)
    return env_path or r"C:\Scripts\chromedriver.exe"

CHROMEDRIVER_PATH = _resolve_chromedriver()

# Configure Logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger()

def _find_window_containing(substr):
    """Chrome appends ' - Google Chrome' (or similar) to the page title, so
    an exact FindWindowW match on our marker always fails - enumerate
    top-level windows and match by substring instead."""
    user32 = ctypes.windll.user32
    found = []
    EnumWindowsProc = ctypes.WINFUNCTYPE(ctypes.c_bool, ctypes.c_int, ctypes.c_int)

    def _callback(hwnd, _lparam):
        length = user32.GetWindowTextLengthW(hwnd)
        if length:
            buf = ctypes.create_unicode_buffer(length + 1)
            user32.GetWindowTextW(hwnd, buf, length + 1)
            if substr in buf.value:
                found.append(hwnd)
        return True

    user32.EnumWindows(EnumWindowsProc(_callback), 0)
    return found[0] if found else 0

def _restore_title(driver, title):
    """Put a previously-saved document.title back (no-op if none was saved).
    Best-effort: a failure here must never abort the run."""
    if title is None:
        return
    try:
        driver.execute_script("document.title = arguments[0];", title)
    except Exception:
        pass

def bring_chrome_to_foreground(driver):
    """Force the Selenium-controlled Chrome window to the OS foreground.
    Windows blocks background processes from stealing real input focus by
    default (the exact situation here: PS Automation itself runs hidden),
    so this combines two techniques for reliability: (1) SetWindowPos with
    HWND_TOPMOST, which forces the window to the top of the Z-order
    regardless of focus rules - this alone is enough for screenshots, since
    they just capture whatever is visually on top; (2) AttachThreadInput,
    which lets SetForegroundWindow succeed even while another window
    currently holds real/recent user input. Tags the page with a unique
    title first so the right window is found even with other Chrome
    windows open."""
    original_title = None
    try:
        # Remember the page's REAL title before tagging it, so we can put it
        # back the instant the window is located - the marker must never
        # linger in the browser tab captured in screenshots.
        try:
            original_title = driver.title
        except Exception:
            original_title = None

        marker = f"PSAUTO_FOCUS_{os.getpid()}"
        driver.execute_script(f"document.title = '{marker}';")
        user32 = ctypes.windll.user32
        kernel32 = ctypes.windll.kernel32
        hwnd = 0
        for _ in range(20):
            hwnd = _find_window_containing(marker)
            if hwnd:
                break
            time.sleep(0.15)

        # Marker has served its purpose (or we gave up) - restore the real
        # title now, before any screenshot runs.
        _restore_title(driver, original_title)
        original_title = None  # already restored; don't restore again in finally

        if not hwnd:
            logger.warning("Could not locate the Chrome window to bring it to the foreground.")
            return

        SW_SHOWMAXIMIZED, HWND_TOPMOST, HWND_NOTOPMOST = 3, -1, -2
        SWP_NOMOVE, SWP_NOSIZE, SWP_SHOWWINDOW = 0x0002, 0x0001, 0x0040

        # Show the window MAXIMIZED - NOT SW_RESTORE, which would un-maximize the
        # window that --start-maximized / maximize_window() just maximized,
        # leaving a small windowed Chrome captured in the screenshots.
        user32.ShowWindow(hwnd, SW_SHOWMAXIMIZED)
        user32.SetWindowPos(hwnd, HWND_TOPMOST, 0, 0, 0, 0, SWP_NOMOVE | SWP_NOSIZE | SWP_SHOWWINDOW)

        current_tid = kernel32.GetCurrentThreadId()
        fg_hwnd = user32.GetForegroundWindow()
        fg_tid = user32.GetWindowThreadProcessId(fg_hwnd, None) if fg_hwnd else 0
        target_tid = user32.GetWindowThreadProcessId(hwnd, None)

        attached_fg = attached_target = False
        if fg_tid and fg_tid != current_tid:
            attached_fg = bool(user32.AttachThreadInput(current_tid, fg_tid, True))
        if target_tid and target_tid != current_tid:
            attached_target = bool(user32.AttachThreadInput(current_tid, target_tid, True))
        try:
            user32.BringWindowToTop(hwnd)
            user32.SetForegroundWindow(hwnd)
        finally:
            if attached_fg:
                user32.AttachThreadInput(current_tid, fg_tid, False)
            if attached_target:
                user32.AttachThreadInput(current_tid, target_tid, False)

        user32.SetWindowPos(hwnd, HWND_NOTOPMOST, 0, 0, 0, 0, SWP_NOMOVE | SWP_NOSIZE)
    except Exception as e:
        logger.warning(f"Could not bring Chrome window to foreground: {e}")
    finally:
        # If we bailed out before the mid-function restore, make sure the
        # marker still never survives this call.
        _restore_title(driver, original_title)

def get_screen_resolution():
    """Gets the resolution of the primary monitor only with DPI awareness."""
    try:
        user32 = ctypes.windll.user32
        user32.SetProcessDPIAware()
        return user32.GetSystemMetrics(0), user32.GetSystemMetrics(1)
    except Exception as e:
        logger.error(f"Error getting screen resolution: {e}")
        return 1920, 1080 

def take_full_screen_shot(filename):
    """Captures the primary monitor only (ignores extended displays)."""
    try:
        time.sleep(1)
        width, height = get_screen_resolution()
        screenshot = ImageGrab.grab(bbox=(0, 0, width, height))
        screenshot.save(filename)
        return True
    except Exception as e:
        logger.error(f"Error taking full screenshot: {e}")
        return False

def add_to_word(doc, title, image_path):
    """Adds a title, timestamp, and image to the Word document."""
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_heading(title, level=2)
    doc.add_paragraph(f"Taken at: {now}") 
    try:
        doc.add_picture(image_path, width=Inches(6.5))
    except Exception as e:
        doc.add_paragraph(f"Error adding image: {e}")
    doc.add_paragraph("\n")

def send_key_sequence(driver, key, delay=1.0):
    """Sends a specific key to the browser with a delay."""
    try:
        actions = ActionChains(driver)
        actions.send_keys(key).perform()
        time.sleep(delay)
    except Exception as e:
        logger.error(f"Error sending key: {e}")

# US-keyboard shifted symbols, so an arbitrary password/username can be typed
# into the DCUI console applet the same reliable way the old hardcoded value
# was (explicit SHIFT key-down/up around each shifted character) - the applet
# needs the real shift events, a plain send_keys of e.g. '!' or 'C' doesn't
# register correctly.
_SHIFT_SYMBOLS = {
    '!': '1', '@': '2', '#': '3', '$': '4', '%': '5', '^': '6', '&': '7',
    '*': '8', '(': '9', ')': '0', '_': '-', '+': '=', '{': '[', '}': ']',
    '|': '\\', ':': ';', '"': "'", '<': ',', '>': '.', '?': '/', '~': '`',
}

def type_secure_string(driver, text):
    """Type an arbitrary string (username/password supplied via env, never
    hardcoded) into the focused DCUI field, applying a real SHIFT for
    uppercase letters and shifted symbols so every character registers."""
    actions = ActionChains(driver)
    actions.reset_actions()
    for ch in text:
        if ch.isalpha() and ch.isupper():
            actions.key_down(Keys.SHIFT).pause(0.15).send_keys(ch.lower()).pause(0.15).key_up(Keys.SHIFT).pause(0.15)
        elif ch in _SHIFT_SYMBOLS:
            actions.key_down(Keys.SHIFT).pause(0.15).send_keys(_SHIFT_SYMBOLS[ch]).pause(0.15).key_up(Keys.SHIFT).pause(0.15)
        else:
            actions.send_keys(ch).pause(0.2)
    actions.perform()

def get_hostname_via_local_racadm(ip_address):
    """
    מריץ את פקודת racadm המקומית של מערכת ההפעלה ומבצע שאילתה מרחוק 
    מול ה-IP הספציפי עם המשתמש והסיסמה שהוגדרו.
    """
    logger.info(f"Executing local RACADM CLI for IP: {ip_address}...")
    
    cmd = ["racadm", "-r", ip_address, "-u", USERNAME, "-p", PASSWORD, "get", "iDRAC.NIC.DNSRacName"]
    
    try:
        process = subprocess.Popen(cmd, shell=False, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        stdout, stderr = process.communicate()
        output = stdout.decode('utf-8', errors='ignore').strip()
        
        if output and "ERROR:" not in output:
            if "=" in output:
                output = output.split("=")[-1].strip()
            
            final_name = output.split('\n')[0].strip()
            if final_name:
                logger.info(f"SUCCESS! Local RACADM Extracted Hostname: {final_name}")
                return final_name
                
    except Exception as e:
        logger.error(f"Failed to execute local racadm command: {e}")
        
    return "Unknown_Host"

def process_single_server_dcui(server, doc):
    name = server['name']
    url = server['url']
    clean_ip = url.replace("https://", "").replace("/", "")
    
    logger.info(f"==========================================")
    logger.info(f"Processing Server (DCUI): {name} ({url})")

    # שליפת שם השרת באמצעות ה-RACADM המקומי לפני תחילת תהליך ה-GUI
    extracted_hostname = get_hostname_via_local_racadm(clean_ip)
    logger.info(f"Final Hostname resolved: {extracted_hostname}")

    # הוספת כותרת ראשית ומדויקת לדוח הוורד על בסיס השם שחולץ מה-RACADM
    doc.add_heading(f"Server: {extracted_hostname} ({clean_ip})", level=1)

    options = webdriver.ChromeOptions()
    options.add_argument("--start-maximized")
    options.add_argument("--ignore-certificate-errors")
    options.add_argument("--allow-running-insecure-content")
    options.add_argument("--no-proxy-server") 
    
    # הסרת באנר האוטומציה לחלוטין
    options.add_experimental_option("excludeSwitches", ["enable-automation"])
    options.add_experimental_option('useAutomationExtension', False)
    options.add_argument("--disable-blink-features=AutomationControlled")
    
    service = Service(executable_path=CHROMEDRIVER_PATH)
    driver = webdriver.Chrome(service=service, options=options)

    # Force the window to fill the screen. This is the reliable, canonical way -
    # the --start-maximized flag alone is not always honored, and the screenshots
    # grab the whole primary monitor, so a maximized window is what makes them
    # look right (rather than a small window with desktop around it).
    try:
        driver.maximize_window()
    except Exception:
        pass

    driver.switch_to.window(driver.current_window_handle)
    bring_chrome_to_foreground(driver)

    try:
        driver.set_page_load_timeout(60)
        driver.get(url)
        bring_chrome_to_foreground(driver)  # retry once a real page is loaded

        # --- WEB LOGIN ---
        logger.info("Attempting Web Login...")
        try:
            wait = WebDriverWait(driver, 20)
            
            user_field = wait.until(EC.element_to_be_clickable((By.NAME, "username")))
            user_field.clear()
            user_field.send_keys(USERNAME)
            time.sleep(0.5)

            pass_field = driver.find_element(By.NAME, "password")
            driver.execute_script("arguments[0].value = arguments[1];", pass_field, PASSWORD)
            
            pass_field.send_keys(" ") 
            pass_field.send_keys(Keys.BACKSPACE)
            time.sleep(0.5)
            
            pass_field.send_keys(Keys.RETURN)
            logger.info("Login submitted. Waiting 15s...")
            time.sleep(15)
            
            # סגירת פופאפים במידה וקיימים
            try:
                popups = driver.find_elements(By.XPATH, "//*[contains(text(), 'Later') or contains(text(), 'Keep Default')]")
                for p in popups:
                    if p.is_displayed(): p.click()
            except:
                pass

        except Exception as e:
            logger.error(f"Web Login process error: {e}")
            return "Failed", f"Web Login Error"

        # --- LAUNCH VIRTUAL CONSOLE ---
        logger.info("Launching Virtual Console...")
        main_window = driver.current_window_handle
        
        try:
            console_btn = WebDriverWait(driver, 15).until(
                EC.element_to_be_clickable((By.XPATH, "//*[@id='console_preview_img_id'] | //button[contains(., 'Virtual Console')]"))
            )
            console_btn.click()
            time.sleep(5)

            all_windows = driver.window_handles
            if len(all_windows) > 1:
                for window in all_windows:
                    if window != main_window:
                        driver.switch_to.window(window)
                        driver.maximize_window()
                        logger.info("Switched to Virtual Console window.")
                        break
            else:
                logger.error("No new window detected for console.")
                return "Failed", "Console Window not found"

            logger.info("Waiting 4s for Console connection...")
            time.sleep(4)

            try:
                driver.find_element(By.TAG_NAME, "body").click()
                time.sleep(0.5)
            except:
                pass

            # 1. DCUI DASHBOARD
            logger.info("Sending ENTER before screenshot...")
            send_key_sequence(driver, Keys.ENTER, delay=2) 

            logger.info("Capturing: DCUI Dashboard")
            img_name = f"temp_{name}_1_DCUI_Dashboard.png"
            if take_full_screen_shot(img_name):
                add_to_word(doc, "DCUI Dashboard", img_name)
                os.remove(img_name)

            # DCUI LOGIN (F2)
            logger.info("Sending F2 to Log in...")
            send_key_sequence(driver, Keys.F2, delay=2) 
            send_key_sequence(driver, Keys.F2, delay=2)

            try:
                driver.find_element(By.TAG_NAME, "body").click()
            except:
                pass

            time.sleep(2)

            logger.info("Clearing User field...")
            actions = ActionChains(driver)
            actions.reset_actions()
            for _ in range(20): 
                actions.send_keys(Keys.BACK_SPACE)
                actions.pause(0.05)
            actions.perform()
            time.sleep(0.5)

            # STEP 1: TYPE USERNAME (from env, never hardcoded)
            logger.info("Typing username...")
            type_secure_string(driver, USERNAME)

            time.sleep(1)
            send_key_sequence(driver, Keys.ENTER, delay=2)

            # STEP 2: TYPE PASSWORD (from env, never hardcoded)
            logger.info("Typing password...")
            try:
                driver.find_element(By.TAG_NAME, "body").click()
            except:
                pass

            type_secure_string(driver, PASSWORD)

            time.sleep(2)
            send_key_sequence(driver, Keys.ENTER, delay=4)

            # NAVIGATE TO MANAGEMENT NETWORK
            logger.info("Navigating to Management Network...")
            send_key_sequence(driver, Keys.ARROW_DOWN, delay=0.8)
            send_key_sequence(driver, Keys.ENTER, delay=2)

            # 2. NETWORK ADAPTERS
            logger.info("Capturing: Network Adapters")
            send_key_sequence(driver, Keys.ENTER, delay=2)
            img_name = f"temp_{name}_2_Network_Adapters.png"
            if take_full_screen_shot(img_name):
                add_to_word(doc, "Network Adapters", img_name)
                os.remove(img_name)
            send_key_sequence(driver, Keys.ESCAPE, delay=2)

            # 2b. VLAN CONFIGURATION
            logger.info("Capturing: VLAN Configuration")
            send_key_sequence(driver, Keys.ARROW_DOWN, delay=0.8)
            send_key_sequence(driver, Keys.ENTER, delay=2)
            img_name = f"temp_{name}_2b_VLAN.png"
            if take_full_screen_shot(img_name):
                add_to_word(doc, "VLAN Configuration", img_name)
                os.remove(img_name)
            send_key_sequence(driver, Keys.ESCAPE, delay=2)

            # 3. IPv4 CONFIGURATION
            logger.info("Capturing: IPv4 Configuration")
            send_key_sequence(driver, Keys.ARROW_DOWN, delay=0.8)
            send_key_sequence(driver, Keys.ENTER, delay=2)
            img_name = f"temp_{name}_3_IPv4.png"
            if take_full_screen_shot(img_name):
                add_to_word(doc, "IPv4 Configuration", img_name)
                os.remove(img_name)
            send_key_sequence(driver, Keys.ESCAPE, delay=2)

            # 4. IPv6 CONFIGURATION
            logger.info("Capturing: IPv6 Configuration")
            send_key_sequence(driver, Keys.ARROW_DOWN, delay=0.8)
            send_key_sequence(driver, Keys.ENTER, delay=2)
            img_name = f"temp_{name}_4_IPv6.png"
            if take_full_screen_shot(img_name):
                add_to_word(doc, "IPv6 Configuration", img_name)
                os.remove(img_name)
            send_key_sequence(driver, Keys.ESCAPE, delay=2)

            # 5. DNS CONFIGURATION
            logger.info("Capturing: DNS Configuration")
            send_key_sequence(driver, Keys.ARROW_DOWN, delay=0.8)
            send_key_sequence(driver, Keys.ENTER, delay=2)
            img_name = f"temp_{name}_5_DNS.png"
            if take_full_screen_shot(img_name):
                add_to_word(doc, "DNS Configuration", img_name)
                os.remove(img_name)
            send_key_sequence(driver, Keys.ESCAPE, delay=2)

            # CLEANUP
            logger.info("Exiting DCUI menus...")
            for _ in range(5):
                send_key_sequence(driver, Keys.ESCAPE, delay=0.5)

            driver.close()
            driver.switch_to.window(main_window)

            return "Success", f"Captured ({extracted_hostname})"

        except Exception as e:
            logger.error(f"Console navigation error: {e}")
            return "Failed", f"Console Error"

    except Exception as e:
        logger.error(f"CRITICAL ERROR on {name}: {e}")
        return "Failed", str(e)
    
    finally:
        driver.quit()
        logger.info(f"Finished {name}.")

def main():
    if not os.path.exists(CHROMEDRIVER_PATH):
        logger.error(f"Error: chromedriver.exe not found at {CHROMEDRIVER_PATH}")
        return

    print("\n" + "="*50)
    print("   iDRAC DCUI Screenshot Tool (Local RACADM Integration)")
    print("="*50)
    print("Select Mode:")
    print("1. Sequence Mode (Base IP + Start Suffix + Count)")
    print("2. Specific IP List Mode (Enter full IPs manually)")
    print("="*50)
    
    mode_choice = input("Enter choice (1 or 2): ").strip()
    servers_list = []
    
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    dynamic_output_file = "ESX_DCUI_Report.docx"

    if mode_choice == "1":
        base_ip_prefix = input("Enter Base IP (e.g., 10.201.91): ").strip()
        try:
            start_suffix = int(input("Enter Starting Suffix (e.g., 207): "))
            server_count = int(input("Enter number of servers to scan (e.g., 9): "))
        except ValueError:
            print("Error: Suffix and Count must be numbers.")
            return

        dynamic_output_file = f"ESX_DCUI_Report_{base_ip_prefix}.{start_suffix}_{timestamp}.docx"
        for i in range(server_count):
            current_suffix = start_suffix + i
            ip = f"{base_ip_prefix}.{current_suffix}"
            servers_list.append({"name": f"ESX{i + 1}", "url": f"https://{ip}/"})

    elif mode_choice == "2":
        print("\n--- Mode 2 Selected ---")
        counter = 1
        while True:
            user_ip = input(f"Enter IP #{counter}: ").strip()
            if user_ip.lower() == "done" or user_ip == "": break
            if len(user_ip) < 7: continue
            servers_list.append({"name": f"ESX{counter} ({user_ip})", "url": f"https://{user_ip}/"})
            counter += 1
        
        if not servers_list: return
        first_ip = servers_list[0]['url'].replace("https://", "").replace("/", "")
        dynamic_output_file = f"ESX_DCUI_Report_{first_ip}_CustomList_{timestamp}.docx"
    else:
        print("Invalid choice.")
        return

    print("-" * 50)
    print(f"Starting DCUI Process for {len(servers_list)} servers...")
    print(f"Output File: {dynamic_output_file}")
    print("-" * 50)

    doc = Document()
    doc.add_heading('iDRAC DCUI Report', 0)
    execution_summary = []

    def safe_save():
        try:
            doc.save(dynamic_output_file)
            logger.info(f"Progress saved incrementally to: {dynamic_output_file}")
        except PermissionError:
            logger.error(f"ERROR: Could not save output file. Please close '{dynamic_output_file}' if it is open.")

    # מנגנון הגנה לשמירת הקובץ בכל מקרה של עצירה/קריסה - שומר אחרי כל שרת
    # בנפרד (לא רק ב-finally), כי עצירה בכוח (Kill) לא מפעילה finally כלל.
    try:
        for server in servers_list:
            status, msg = process_single_server_dcui(server, doc)
            execution_summary.append({
                "name": server['name'],
                "url": server['url'],
                "status": status,
                "msg": msg
            })
            safe_save()
    except (KeyboardInterrupt, Exception) as master_error:
        logger.warning(f"Process interrupted! Saving all data captured so far. Info: {master_error}")
    finally:
        safe_save()
        logger.info(f"SUCCESS: Safe-save executed. Report successfully saved to: {dynamic_output_file}")

    print("\n\n" + "="*80)
    print(f"{'FINAL EXECUTION SUMMARY (DCUI)':^80}")
    print("="*80)
    for item in execution_summary:
        clean_ip = item['url'].replace("https://", "").replace("/", "")
        print(f"{item['name']:<20} | {clean_ip:<25} | {item['status']:<10} | {item['msg']}")
    print("="*80 + "\n")

if __name__ == "__main__":
    main()