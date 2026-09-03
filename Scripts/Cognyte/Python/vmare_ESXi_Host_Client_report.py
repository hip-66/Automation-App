import sys
import subprocess
import os
import time
import datetime
import logging
import ctypes

# --- AUTO-INSTALL DEPENDENCIES ---
def install_dependencies():
    """Checks for required libraries and installs them if missing."""
    required_packages = ["selenium", "python-docx", "Pillow"]
    for package in required_packages:
        try:
            if package == "python-docx": import docx
            elif package == "Pillow": import PIL
            else: __import__(package.replace("-", "_"))
        except ImportError:
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])

install_dependencies()

from docx import Document
from docx.shared import Inches
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from PIL import ImageGrab

# --- Configuration ---
# Never hardcoded: PSAUTO_USERNAME/PASSWORD (explicit override from the app's
# UI) wins; otherwise PSAUTO_DEFAULT_USERNAME/PASSWORD (the app's encrypted
# .env default) is used; a standalone run with neither set prompts instead.
USERNAME = os.environ.get("PSAUTO_USERNAME", "").strip() or os.environ.get("PSAUTO_DEFAULT_USERNAME", "").strip()
PASSWORD = os.environ.get("PSAUTO_PASSWORD", "") or os.environ.get("PSAUTO_DEFAULT_PASSWORD", "")
if not USERNAME or not PASSWORD:
    if sys.stdin is not None and sys.stdin.isatty():
        import getpass
        if not USERNAME:
            USERNAME = input("ESXi username: ").strip()
        if not PASSWORD:
            PASSWORD = getpass.getpass("ESXi password: ")
    else:
        print("ERROR: No ESXi username/password available (defaults not configured) and this run "
              "is non-interactive, so a prompt cannot be shown. Configure the app's default "
              "ESXi credentials and retry.", flush=True)
        sys.exit(1)

def _resolve_chromedriver():
    """Prefer the exact driver handed over by the PS Automation app (env var).
    Otherwise search upward for the app's Chromedrivers folder and pick the
    newest version in it - so manual runs also get a sane driver."""
    env_path = os.environ.get("CHROMEDRIVER_PATH")
    if env_path and os.path.exists(env_path):
        return env_path
    search_dir = os.path.dirname(os.path.abspath(__file__))
    for _ in range(6):
        drivers_root = os.path.join(search_dir, "Chromedrivers")
        if os.path.isdir(drivers_root):
            import re
            candidates = []
            for root, _dirs, files in os.walk(drivers_root):
                for fname in files:
                    if fname.lower().startswith("chromedriver") and fname.lower().endswith(".exe"):
                        rel = os.path.relpath(os.path.join(root, fname), drivers_root)
                        m = re.search(r"(\d+)", rel)
                        candidates.append((int(m.group(1)) if m else -1, os.path.join(root, fname)))
            if candidates:
                candidates.sort(key=lambda c: c[0], reverse=True)
                return candidates[0][1]
        search_dir = os.path.dirname(search_dir)
    return env_path or r"C:\Scripts\chromedriver.exe"

CHROMEDRIVER_PATH = _resolve_chromedriver()

# Configure Logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger()

def _find_window_containing(substr):
    """Chrome appends ' - Google Chrome' (or similar) to the page title, so
    an exact FindWindowW match on our marker always fails - enumerate
    top-level windows and match by substring instead."""
    user32 = ctypes.windll.user32
    found = []
    EnumWindowsProc = ctypes.WINFUNCTYPE(ctypes.c_bool, ctypes.c_int, ctypes.c_int)

    def _callback(hwnd, _lparam):
        length = user32.GetWindowTextLengthW(hwnd)
        if length:
            buf = ctypes.create_unicode_buffer(length + 1)
            user32.GetWindowTextW(hwnd, buf, length + 1)
            if substr in buf.value:
                found.append(hwnd)
        return True

    user32.EnumWindows(EnumWindowsProc(_callback), 0)
    return found[0] if found else 0

def _restore_title(driver, title):
    """Put a previously-saved document.title back (no-op if none was saved).
    Best-effort: a failure here must never abort the run."""
    if title is None:
        return
    try:
        driver.execute_script("document.title = arguments[0];", title)
    except Exception:
        pass

def bring_chrome_to_foreground(driver):
    """Force the Selenium-controlled Chrome window to the OS foreground.
    Windows blocks background processes from stealing real input focus by
    default (the exact situation here: PS Automation itself runs hidden),
    so this combines two techniques for reliability: (1) SetWindowPos with
    HWND_TOPMOST, which forces the window to the top of the Z-order
    regardless of focus rules - this alone is enough for screenshots, since
    they just capture whatever is visually on top; (2) AttachThreadInput,
    which lets SetForegroundWindow succeed even while another window
    currently holds real/recent user input. Tags the page with a unique
    title first so the right window is found even with other Chrome
    windows open."""
    original_title = None
    try:
        # Remember the page's REAL title before tagging it, so we can put it
        # back the instant the window is located - the marker must never
        # linger in the browser tab captured in screenshots.
        try:
            original_title = driver.title
        except Exception:
            original_title = None

        marker = f"PSAUTO_FOCUS_{os.getpid()}"
        driver.execute_script(f"document.title = '{marker}';")
        user32 = ctypes.windll.user32
        kernel32 = ctypes.windll.kernel32
        hwnd = 0
        for _ in range(20):
            hwnd = _find_window_containing(marker)
            if hwnd:
                break
            time.sleep(0.15)

        # Marker has served its purpose (or we gave up) - restore the real
        # title now, before any screenshot runs.
        _restore_title(driver, original_title)
        original_title = None  # already restored; don't restore again in finally

        if not hwnd:
            logger.warning("Could not locate the Chrome window to bring it to the foreground.")
            return

        SW_SHOWMAXIMIZED, HWND_TOPMOST, HWND_NOTOPMOST = 3, -1, -2
        SWP_NOMOVE, SWP_NOSIZE, SWP_SHOWWINDOW = 0x0002, 0x0001, 0x0040

        # Show the window MAXIMIZED - NOT SW_RESTORE, which would un-maximize the
        # window that --start-maximized / maximize_window() just maximized,
        # leaving a small windowed Chrome captured in the screenshots.
        user32.ShowWindow(hwnd, SW_SHOWMAXIMIZED)
        user32.SetWindowPos(hwnd, HWND_TOPMOST, 0, 0, 0, 0, SWP_NOMOVE | SWP_NOSIZE | SWP_SHOWWINDOW)

        current_tid = kernel32.GetCurrentThreadId()
        fg_hwnd = user32.GetForegroundWindow()
        fg_tid = user32.GetWindowThreadProcessId(fg_hwnd, None) if fg_hwnd else 0
        target_tid = user32.GetWindowThreadProcessId(hwnd, None)

        attached_fg = attached_target = False
        if fg_tid and fg_tid != current_tid:
            attached_fg = bool(user32.AttachThreadInput(current_tid, fg_tid, True))
        if target_tid and target_tid != current_tid:
            attached_target = bool(user32.AttachThreadInput(current_tid, target_tid, True))
        try:
            user32.BringWindowToTop(hwnd)
            user32.SetForegroundWindow(hwnd)
        finally:
            if attached_fg:
                user32.AttachThreadInput(current_tid, fg_tid, False)
            if attached_target:
                user32.AttachThreadInput(current_tid, target_tid, False)

        user32.SetWindowPos(hwnd, HWND_NOTOPMOST, 0, 0, 0, 0, SWP_NOMOVE | SWP_NOSIZE)
    except Exception as e:
        logger.warning(f"Could not bring Chrome window to foreground: {e}")
    finally:
        # If we bailed out before the mid-function restore, make sure the
        # marker still never survives this call.
        _restore_title(driver, original_title)

def get_screen_resolution():
    """Gets the resolution of the primary monitor only (DPI aware)."""
    user32 = ctypes.windll.user32
    user32.SetProcessDPIAware()
    return user32.GetSystemMetrics(0), user32.GetSystemMetrics(1)

def take_primary_monitor_screenshot(filename):
    """Captures ONLY the primary monitor to avoid multi-screen issues."""
    try:
        time.sleep(2) # Stabilization wait
        width, height = get_screen_resolution()
        screenshot = ImageGrab.grab(bbox=(0, 0, width, height))
        screenshot.save(filename)
        return True
    except Exception as e:
        logger.error(f"Screenshot error: {e}")
        return False

def add_to_word(doc, title, image_path, level=2):
    """Adds title, timestamp and image to the Word document."""
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_heading(title, level=level)
    doc.add_paragraph(f"Captured at: {now}")
    try:
        doc.add_picture(image_path, width=Inches(6.5))
    except Exception as e:
        doc.add_paragraph(f"Error adding image: {e}")
    doc.add_paragraph("\n")

def esxi_force_click(driver, text):
    """Finds an element by its exact text content and clicks it via JS."""
    script = f"""
    var elements = document.querySelectorAll('.vui-label, .nav-link, button, a, span, .vui-tab');
    for (var i = 0; i < elements.length; i++) {{
        if (elements[i].textContent.trim().toLowerCase() === '{text.lower()}') {{
            elements[i].click();
            return true;
        }}
    }}
    return false;
    """
    success = driver.execute_script(script)
    if success:
        time.sleep(3) # Wait for UI transition
    return success

def process_esxi_server(ip, doc):
    """Handles logic for a single host: Login, navigation through all tags, and screenshots."""
    url = f"https://{ip}/ui/#/login"
    logger.info(f"Starting process for: {ip}")
    
    # Structure the document headers by server
    doc.add_heading(f"Server Host: {ip}", level=1)
    
    options = webdriver.ChromeOptions()
    options.add_argument("--start-maximized")
    options.add_argument("--ignore-certificate-errors")
    options.add_argument("--no-proxy-server")
    
    # --- הגדרות מתקדמות ומקיפות להעלמת באנר האוטומציה לחלוטין ---
    options.add_argument("--disable-infobars")
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.exclude_switches = ["enable-automation"]
    options.add_experimental_option('useAutomationExtension', False)
    options.add_experimental_option("excludeSwitches", ["enable-automation"])
    
    service = Service(executable_path=CHROMEDRIVER_PATH)
    driver = webdriver.Chrome(service=service, options=options)

    # Force the window to fill the screen. This is the reliable, canonical way -
    # the --start-maximized flag alone is not always honored, and the screenshots
    # grab the whole primary monitor, so a maximized window is what makes them
    # look right (rather than a small window with desktop around it).
    try:
        driver.maximize_window()
    except Exception:
        pass

    bring_chrome_to_foreground(driver)

    # חסימת ה-Flag של האוטומציה גם דרך הרצת קוד בדפדפן
    driver.execute_cdp_cmd("Page.addScriptToEvaluateOnNewDocument", {
        "source": "Object.defineProperty(navigator, 'webdriver', {get: () => undefined})"
    })
    
    try:
        driver.get(url)
        bring_chrome_to_foreground(driver)  # retry once a real page is loaded
        wait = WebDriverWait(driver, 30)
        
        # --- LOGIN ---
        wait.until(EC.element_to_be_clickable((By.ID, "username"))).send_keys(USERNAME)
        driver.find_element(By.ID, "password").send_keys(PASSWORD + Keys.ENTER)
        
        # Wait for UI to load fully
        time.sleep(10)
        
        # Minimize 'Recent Tasks' bar
        driver.execute_script("try { document.querySelector('em[title=\"Minimize\"]').click(); } catch(e) {}")
        
        # --- צילום DASHBOARD עם הקטנת מסך ל-80% ---
        logger.info("Zooming out to 80% for Dashboard screenshot...")
        driver.execute_script("document.body.style.zoom='0.8'")
        time.sleep(2)
        
        img = f"host_{ip}.png"
        if take_primary_monitor_screenshot(img):
            add_to_word(doc, f"Host Summary - {ip}", img)
            os.remove(img)
            
        # החזרה לגודל המקורי (CTRL+0)
        logger.info("Restoring zoom to 100%...")
        driver.execute_script("document.body.style.zoom='1.0'")
        time.sleep(1)

        # --- NAVIGATION SEQUENCE ---
        
        # FIRST STEP: Navigate to Manage -> Time & date immediately
        logger.info("Navigating straight into Manage -> Time & date view...")
        esxi_force_click(driver, "Manage")
        time.sleep(3)
        
        try:
            wait.until(EC.presence_of_element_located((By.CLASS_NAME, "vui-tertiary-tabs")))
            time_date_xpath = "//div[contains(@class,'vui-tertiary-tabs')]//a[contains(text(),'Time & date')] | //a[contains(text(),'Time & date')]"
            time_date_link = wait.until(EC.presence_of_element_located((By.XPATH, time_date_xpath)))
            
            driver.execute_script("""
                var el = arguments[0];
                el.scrollIntoView({block: "center", inline: "nearest"});
                var actions = ['focus', 'mousedown', 'mouseup', 'click'];
                for (var i = 0; i < actions.length; i++) {
                    var ev = new MouseEvent(actions[i], {
                        bubbles: true,
                        cancelable: true,
                        view: window
                    });
                    el.dispatchEvent(ev);
                }
            """, time_date_link)
            logger.info("Successfully executed event sequence on Time & date tab.")
        except Exception as e:
            logger.warning(f"Primary element target failed, trying textual DOM script fallback... Error: {e}")
            driver.execute_script("""
                var links = document.querySelectorAll('a, .vui-tab, span, li');
                for (var i = 0; i < links.length; i++) {
                    if (links[i].textContent.trim() === 'Time & date') {
                        links[i].click();
                        break;
                    }
                }
            """)
        
        time.sleep(6)
        img = "manage_time_date.png"
        if take_primary_monitor_screenshot(img):
            add_to_word(doc, f"Manage Time & Date - {ip}", img)
            os.remove(img)

        # SECOND STEP: Move into Manage -> Licensing
        logger.info("Proceeding to Manage -> Licensing...")
        esxi_force_click(driver, "Manage")
        esxi_force_click(driver, "Licensing")
        img = "manage_licensing.png"
        if take_primary_monitor_screenshot(img):
            add_to_word(doc, f"Manage Licensing - {ip}", img)
            os.remove(img)

        # --- מיקום חדש: צילום VIRTUAL MACHINES (מיד לאחר סיום הרישיון) ---
        logger.info("Navigating to Virtual Machines...")
        esxi_force_click(driver, "Virtual Machines")
        time.sleep(4)
        img = "virtual_machines.png"
        if take_primary_monitor_screenshot(img):
            add_to_word(doc, f"Virtual Machines - {ip}", img)
            os.remove(img)
            
        # Click into EVERY Virtual Machine in turn to photograph each one's
        # details screen (same discover-all/loop/navigate-back pattern already
        # used below for Port Groups and Virtual Switches).
        try:
            vm_links = driver.find_elements(By.XPATH, "//table[contains(@class,'vui-grid')]//tbody//tr//td[2]//a | //table//tbody//tr//td[1]//a | //a[contains(@href, '#/host/vms/')]")
            vm_names = [elem.text.strip() for elem in vm_links if elem.text.strip()]

            logger.info(f"Discovered Virtual Machines dynamically: {vm_names}")

            for vm_name in vm_names:
                logger.info(f"Navigating into Virtual Machine: {vm_name}")

                click_script = f"""
                    var links = document.querySelectorAll('table tbody tr td a, a[href*="#/host/vms/"]');
                    for (var i = 0; i < links.length; i++) {{
                        if (links[i].textContent.trim() === '{vm_name}') {{
                            links[i].click();
                            return true;
                        }}
                    }}
                    return false;
                """
                driver.execute_script(click_script)
                time.sleep(5)

                img_vm = f"vm_{vm_name}.png"
                if take_primary_monitor_screenshot(img_vm):
                    add_to_word(doc, f"Virtual Machine Details ({vm_name}) - {ip}", img_vm)
                    os.remove(img_vm)

                # Navigate back to the VM list to continue cleanly to the next one
                esxi_force_click(driver, "Virtual Machines")
                time.sleep(3)
        except Exception as e:
            logger.warning(f"Failed dynamically to click and capture VM details: {e}")

        # Storage -> Datastores
        esxi_force_click(driver, "Storage")
        esxi_force_click(driver, "Datastores")
        img = "storage_datastores.png"
        if take_primary_monitor_screenshot(img):
            add_to_word(doc, f"Storage Datastores - {ip}", img)
            os.remove(img)

        # Storage -> Adapters
        esxi_force_click(driver, "Adapters")
        img = "storage_adapters.png"
        if take_primary_monitor_screenshot(img):
            add_to_word(doc, f"Storage Adapters - {ip}", img)
            os.remove(img)

        # Networking -> Port Groups
        esxi_force_click(driver, "Networking")
        esxi_force_click(driver, "Port groups")
        img = "networking_portgroups.png"
        if take_primary_monitor_screenshot(img):
            add_to_word(doc, f"Networking Port Groups - {ip}", img)
            os.remove(img)

        # --- לולאה גנרית ודינמית לכניסה וצילום של כל PORT GROUP לפי הכמות הקיימת ---
        try:
            # שליפת האלמנטים מהטבלה בזמן אמת (יתאים לכל כמות שמות שקיימת בטבלה)
            pg_elements = driver.find_elements(By.XPATH, "//table[contains(@class,'vui-grid')]//tbody//tr//td[1]//a | //table//tbody//tr//td[1]//a")
            pg_names = [elem.text.strip() for elem in pg_elements if elem.text.strip()]
            
            logger.info(f"Discovered Port Groups dynamically: {pg_names}")
            
            for pg_name in pg_names:
                logger.info(f"Navigating into Port Group: {pg_name}")
                
                click_script = f"""
                    var links = document.querySelectorAll('table tbody tr td a');
                    for (var i = 0; i < links.length; i++) {{
                        if (links[i].textContent.trim() === '{pg_name}') {{
                            links[i].click();
                            return true;
                        }}
                    }}
                    return false;
                """
                driver.execute_script(click_script)
                time.sleep(4) # המתנה לטעינת הנתונים הייחודיים של ה-Port Group
                
                img = f"portgroup_{pg_name}_inside.png"
                if take_primary_monitor_screenshot(img):
                    add_to_word(doc, f"Port Group Details - {pg_name} - {ip}", img)
                    os.remove(img)
                
                # חזרה בטוחה לתפריט הראשי של ה-Port Groups לצורך המשך הלולאה המקורית
                esxi_force_click(driver, "Networking")
                esxi_force_click(driver, "Port groups")
                time.sleep(3)
        except Exception as e:
            logger.warning(f"Generic iteration over Port Groups failed: {e}")

        # Networking -> Virtual Switches
        esxi_force_click(driver, "Networking")
        esxi_force_click(driver, "Virtual switches")
        time.sleep(3)
        img = "networking_vswitches.png"
        if take_primary_monitor_screenshot(img):
            add_to_word(doc, f"Networking Virtual Switches - {ip}", img)
            os.remove(img)

        # Dynamic and Generic Virtual Switches Loop
        try:
            switch_elements = driver.find_elements(By.XPATH, "//table[contains(@class,'vui-grid')]//tbody//tr//td[1]//a | //table//tbody//tr//td[1]//a")
            switch_names = [elem.text.strip() for elem in switch_elements if elem.text.strip()]
            
            logger.info(f"Discovered Virtual Switches dynamically: {switch_names}")
            
            for switch_name in switch_names:
                logger.info(f"Navigating into Switch: {switch_name}")
                
                click_script = f"""
                    var links = document.querySelectorAll('table tbody tr td a');
                    for (var i = 0; i < links.length; i++) {{
                        if (links[i].textContent.trim() === '{switch_name}') {{
                            links[i].click();
                            return true;
                        }}
                    }}
                    return false;
                """
                driver.execute_script(click_script)
                time.sleep(4)
                
                img = f"{switch_name}_inside.png"
                if take_primary_monitor_screenshot(img):
                    add_to_word(doc, f"Virtual Switch Details - {switch_name} - {ip}", img)
                    os.remove(img)
                
                esxi_force_click(driver, "Networking")
                esxi_force_click(driver, "Virtual switches")
                time.sleep(3)
        except Exception as e:
            logger.warning(f"Generic iteration over Virtual Switches failed: {e}")

        # Go back to Networking category to access NICs
        esxi_force_click(driver, "Networking")

        # Networking -> Physical NICs
        esxi_force_click(driver, "Physical NICs")
        img = "physical_nics.png"
        if take_primary_monitor_screenshot(img):
            add_to_word(doc, f"Networking Physical NICs - {ip}", img)
            os.remove(img)

        # Networking -> VMkernel NICs
        esxi_force_click(driver, "VMkernel NICs")
        img = "vmkernel_nics.png"
        if take_primary_monitor_screenshot(img):
            add_to_word(doc, f"Networking VMkernel NICs - {ip}", img)
            os.remove(img)

        driver.quit()
        return "Success"
    except Exception as e:
        logger.error(f"Error on {ip}: {e}")
        driver.quit()
        return "Failed"

def main():
    while True:
        print("\nSelect Mode:")
        print("1. Sequence Mode (Base IP + Start Suffix + Count)")
        print("2. Specific IP List Mode (Enter full IPs manually)")
        print("3. Exit")
        print("\n==================================================\n")
        
        choice = input("Enter choice (1, 2 or 3): ").strip()
        ips = []

        if choice == "3": sys.exit()
        elif choice == "1":
            base = input("Enter Base IP (e.g., 192.168.1): ").strip()
            try:
                start = int(input("Enter Start Suffix: "))
                count = int(input("Enter Count: "))
                ips = [f"{base}.{i}" for i in range(start, start + count)]
            except: continue
        elif choice == "2":
            print("Enter IPs (type 'done' to finish):")
            while True:
                ip_val = input("> ").strip()
                if ip_val.lower() == 'done' or not ip_val: break
                ips.append(ip_val)
        else: continue

        if ips:
            doc = Document()
            doc.add_heading('VMWARE ESXi Host Client Report', 0)
            results = []
            
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"VMWARE_ESXi_Host_Client_{ts}.docx"
            
            try:
                for ip in ips:
                    status = process_esxi_server(ip, doc)
                    results.append((ip, status))
                    doc.save(filename)
                    logger.info(f"Progress saved incrementally to {filename}")
            except KeyboardInterrupt:
                logger.warning("Execution interrupted by user. Saving gathered records...")
            finally:
                doc.save(filename)
                logger.info(f"Emergency final backup successfully written to: {filename}")
            
            print("\n--- FINAL SUMMARY ---")
            for ip, stat in results:
                print(f"{ip}: {stat}")
            print(f"\nReport saved: {filename}")
            
            time.sleep(3)
            sys.exit()

if __name__ == "__main__":
    main()