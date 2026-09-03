# ==============================================================================
# Script Name: iDRAC10_Report.py
# Target: iDRAC 10 Servers ONLY
# Description: Automated script to connect to iDRAC 10 servers, navigate through
#              specific configuration pages, capture screenshots, and compile a 
#              Word document report. Features include precise window manipulation
#              and fault-tolerant element clicking.
# Environment: Offline
# ==============================================================================

import os
import sys
import subprocess
import time
import ctypes
from ctypes import wintypes
from datetime import datetime

# ==============================================================================
# USER INPUTS: Collect metadata for the final Word document.
# ==============================================================================
print("="*50)
print(" iDRAC 10 ATP Report Generator ")
print("="*50)
PO_INPUT = input("Please enter PO: ").strip()  # Purchase Order Number
SO_INPUT = input("Please enter SO: ").strip()  # Sales Order Number
SN_INPUT = input("Please enter SN: ").strip()  # System Serial Number
print("="*50)

# ==============================================================================
# OFFLINE PACKAGE INSTALLATION
# Ensures required Python libraries are installed from a local directory 
# without needing an active internet connection.
# ==============================================================================
OFFLINE_PACKAGES_PATH = r"G:\Automation Library No Internet\offline_packages"

def install_package_offline(package_name, import_name):
    """
    Checks if a Python module is available. If not, installs it using pip 
    from a specified offline local directory.
    """
    try:
        __import__(import_name)
    except ImportError:
        print(f"[INFO] Installing '{package_name}'...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "--no-index", f"--find-links={OFFLINE_PACKAGES_PATH}", package_name])
        except subprocess.CalledProcessError as e:
            print(f"[ERROR] Failed to install '{package_name}'.")
            sys.exit(1)

# Verify and install dependencies
install_package_offline("python-docx", "docx")    # For creating Word documents
install_package_offline("playwright", "playwright") # For browser automation
install_package_offline("Pillow", "PIL")          # For capturing screenshots

import docx
from docx.shared import Inches, Pt
from playwright.sync_api import sync_playwright
from PIL import ImageGrab

# ==============================================================================
# CONFIGURATION VARIABLES
# Defines the target servers and the navigation flow.
# ==============================================================================
SERVERS = {
    "FM1": "192.168.80.122",
    "FM2": "192.168.80.123",
    "PMC1": "192.168.80.124",
    "PMC2": "192.168.80.125",
    "PMC3": "192.168.80.126"
}

USERNAME = "root"
PASSWORD = "admin1234"
CMD_UNIQUE_TITLE = "" # Placeholder for dynamically generated CMD window title

# Defines the exact sequence of clicks to reach target pages and dictating
# whether the page requires a scrolling action to capture all data.
NAVIGATION_PAGES = [
    {"name": "iDRAC view entire server's health status", "click_path": ["Dashboard"], "needs_scroll": False},
    {"name": "Server Names", "click_path": ["iDRAC Settings", "Connectivity", "Network", "Common Settings"], "needs_scroll": False},
    {"name": "IPv4 Settings", "click_path": ["iDRAC Settings", "Connectivity", "Network Interface Settings", "IPv4"], "needs_scroll": False},
    {"name": "Virtual Disks", "click_path": ["Storage", "Overview", "Virtual Disks"], "needs_scroll": False},
    {"name": "Firmware Inventory", "click_path": ["System", "Inventory", "Firmware Inventory"], "needs_scroll": False},
    {"name": "Memory setting", "click_path": ["Configuration", "BIOS Settings", "Memory Settings"], "needs_scroll": False},
    {"name": "Processor setting", "click_path": ["Configuration", "BIOS Settings", "Processor Settings"], "needs_scroll": True},
    {"name": "Boot setting", "click_path": ["Configuration", "BIOS Settings", "Boot Settings"], "needs_scroll": False},
    {"name": "FM Integrated Devices", "click_path": ["Configuration", "BIOS Settings", "More", "Integrated Devices"], "needs_scroll": True},
    {"name": "System profiles settings", "click_path": ["Configuration", "BIOS Settings", "More", "System Profile Settings"], "needs_scroll": True}
]

# ==============================================================================
# OS-LEVEL WINDOW MANIPULATION FUNCTIONS (WINDOWS API)
# ==============================================================================

def get_screen_dimensions():
    """
    Retrieves the primary monitor's resolution (width and height) excluding 
    the taskbar area using Windows API.
    """
    user32 = ctypes.windll.user32
    rect = wintypes.RECT()
    user32.SystemParametersInfoW(48, 0, ctypes.byref(rect), 0) # 48 is SPI_GETWORKAREA
    return rect.right - rect.left, rect.bottom - rect.top

def find_window_by_title_substring(substring):
    """
    Enumerates all open windows to find a specific window handle (HWND) 
    where the window title contains the provided substring.
    """
    hwnd_list = []
    def enum_cb(hwnd, lparam):
        length = ctypes.windll.user32.GetWindowTextLengthW(hwnd)
        if length > 0:
            buff = ctypes.create_unicode_buffer(length + 1)
            ctypes.windll.user32.GetWindowTextW(hwnd, buff, length + 1)
            if substring in buff.value:
                hwnd_list.append(hwnd)
        return True
    ctypes.windll.user32.EnumWindows(ctypes.WINFUNCTYPE(ctypes.c_bool, ctypes.POINTER(ctypes.c_int), ctypes.POINTER(ctypes.c_int))(enum_cb), 0)
    return hwnd_list[0] if hwnd_list else None

def setup_cmd_window(target_x, target_y, target_w, target_h):
    """
    Launches a PowerShell window, retrieves IP config data, and physically 
    manipulates the window (resizing, moving, scrolling) to sit alongside 
    the browser.
    """
    global CMD_UNIQUE_TITLE
    # Create a unique title based on current time to identify the window easily
    CMD_UNIQUE_TITLE = f"iDRAC_CMD_{int(time.time())}"
    
    print("[INFO] Launching PowerShell...")
    # Launch PowerShell with a specific title, widened buffer, and run ipconfig
    ps_command = f'start powershell -NoExit -Command "$Host.UI.RawUI.WindowTitle = \'{CMD_UNIQUE_TITLE}\'; $psw = $Host.UI.RawUI; $sz = $psw.BufferSize; $sz.Width = 200; $psw.BufferSize = $sz; Clear-Host; ipconfig /all"'
    subprocess.Popen(ps_command, shell=True)
    time.sleep(3) # Wait for PowerShell to open and process command
    
    hwnd = find_window_by_title_substring(CMD_UNIQUE_TITLE)
    if not hwnd: return # Exit if window couldn't be found
        
    # Bring the window to the foreground and maximize/restore it
    ctypes.windll.user32.ShowWindow(hwnd, 9)
    ctypes.windll.user32.SetForegroundWindow(hwnd)
    time.sleep(1)
    
    # Virtual Key codes for OS-level keyboard injection
    VK_CONTROL, VK_OEM_MINUS = 0x11, 0xBD
    
    # 1. Inject keyboard commands to Zoom Out (Ctrl + Minus) twice
    print("[INFO] Zooming out CMD (Ctrl + -) twice...")
    for _ in range(2):
        ctypes.windll.user32.keybd_event(VK_CONTROL, 0, 0, 0)
        ctypes.windll.user32.keybd_event(VK_OEM_MINUS, 0, 0, 0)
        time.sleep(0.05)
        ctypes.windll.user32.keybd_event(VK_OEM_MINUS, 0, 2, 0) # Key Up
        ctypes.windll.user32.keybd_event(VK_CONTROL, 0, 2, 0)   # Key Up
        time.sleep(0.3)
        
    # 2. Relocate and resize the CMD window to the right side of the screen
    print("[INFO] Moving CMD to the right side...")
    ctypes.windll.user32.SetWindowPos(hwnd, 0, int(target_x), int(target_y), int(target_w), int(target_h), 0x0040)
    time.sleep(1)
    
    # 3. Physically move the mouse and click the center of the CMD window to force focus
    print("[INFO] Clicking the CMD window to lock focus...")
    click_x, click_y = int(target_x + (target_w / 2)), int(target_y + (target_h / 2))
    ctypes.windll.user32.SetCursorPos(click_x, click_y)
    
    ctypes.windll.user32.mouse_event(0x0002, 0, 0, 0, 0) # Left Mouse Down
    ctypes.windll.user32.mouse_event(0x0004, 0, 0, 0, 0) # Left Mouse Up
    time.sleep(0.5)
    
    # 4. Scroll the CMD window up using simulated mouse wheel events
    print("[INFO] Scrolling CMD to top (Mouse Wheel Scroll)...")
    for _ in range(15): 
        ctypes.windll.user32.mouse_event(0x0800, 0, 0, 120 * 5, 0) # 0x0800 is MOUSEEVENTF_WHEEL
        time.sleep(0.05)

# ==============================================================================
# BROWSER NAVIGATION ENGINE
# ==============================================================================

def smart_click(page, step_text, is_parent_menu=False, child_text=None):
    """
    Intelligently locates and clicks menu items. 
    It checks if parent menus are already expanded before attempting to click them 
    to prevent accidental collapsing.
    """
    # Logic to prevent closing an already opened parent menu
    if is_parent_menu and child_text:
        try:
            # If the child item is visible, the parent is already expanded
            if page.locator(f'text="{child_text}"').first.is_visible(timeout=1000):
                print(f"          [SMART] Parent '{step_text}' is already expanded.")
                return True 
        except: pass
            
    # Attempt 1: Look for exact text match
    try:
        loc = page.get_by_text(step_text, exact=True).filter(visible=True)
        if loc.count() > 0:
            loc.first.click(timeout=3000)
            page.wait_for_timeout(1500) # Wait for animation/load
            return True
    except: pass
        
    # Attempt 2: Look for partial text match (contains) as a fallback
    try:
        loc = page.locator(f"//*[contains(text(), '{step_text}')]").filter(visible=True)
        if loc.count() > 0:
            loc.first.click(timeout=3000)
            page.wait_for_timeout(1500)
            return True
    except: pass
    
    # Return False if the element could not be clicked
    return False

# ==============================================================================
# MAIN AUTOMATION WORKFLOW
# ==============================================================================

def run_automation():
    current_time = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    temp_img_dir = "temp_screenshots" # Directory to temporarily store screenshots
    if not os.path.exists(temp_img_dir): os.makedirs(temp_img_dir)

    # Dictionary to store paths of captured images, organized by page index
    captured_data = {i: {} for i in range(len(NAVIGATION_PAGES))}
    
    # Calculate dimensions: Browser takes 65% width (left), CMD takes 35% (right)
    screen_w, screen_h = get_screen_dimensions()
    browser_w, cmd_w = int(screen_w * 0.65), screen_w - int(screen_w * 0.65)

    try:
        # Prepare the external CMD window
        setup_cmd_window(browser_w, 0, cmd_w, screen_h)

        # Launch the browser using Playwright
        with sync_playwright() as p:
            browser = p.chromium.launch(
                headless=False, # Show the browser UI
                channel="chrome", 
                args=[f'--window-position=0,0', f'--window-size={browser_w},{screen_h}']
            )
            
            for server_name, ip_address in SERVERS.items():
                print(f"\n[INFO] Connecting to Server: {server_name} ({ip_address})")
                context = browser.new_context(ignore_https_errors=True, viewport={'width': browser_w, 'height': screen_h})
                page = context.new_page()
                
                try:
                    # 1. Login Sequence
                    page.goto(f"https://{ip_address}/", wait_until="domcontentloaded", timeout=60000)
                    
                    # Wait for username field to appear
                    user_selector = 'input[name="username"], input[name="user"], #username, #user, #idrac_user'
                    page.wait_for_selector(user_selector, state="visible", timeout=20000)
                    page.wait_for_timeout(2000)
                    
                    # Enter credentials and click Login
                    page.locator(user_selector).first.click() 
                    page.keyboard.type(USERNAME, delay=100) 
                    page.locator('input[name="password"], #password, #idrac_password').first.click()
                    page.keyboard.type(PASSWORD, delay=100)
                    page.locator('button:has-text("Log In"), button:has-text("Login"), #btn-login, button[type="submit"]').first.click()
                    
                    # Wait for Dashboard to load to confirm successful login
                    try: page.wait_for_selector('text="Task Summary"', state="visible", timeout=20000)
                    except: pass
                    page.wait_for_timeout(4000) 

                    # Perform a physical browser zoom equivalent by clicking center and pressing Ctrl + Minus
                    page.mouse.click(browser_w // 2, screen_h // 2)
                    page.keyboard.down('Control')
                    page.keyboard.press('-')
                    page.keyboard.up('Control')
                    page.wait_for_timeout(1000)

                    # 2. Navigation Sequence
                    for page_index, page_info in enumerate(NAVIGATION_PAGES):
                        print(f"\n       -> Navigating to: {page_info['name']}")
                        click_path = page_info['click_path']
                        
                        # Return to the root Dashboard before traversing a new path
                        try:
                            page.locator('text="Dashboard"').first.click(timeout=3000)
                            page.wait_for_timeout(1000)
                        except: pass
                        
                        # Iterate through the defined steps for the current page
                        for i, step in enumerate(click_path):
                            if step == "Dashboard": continue
                            is_parent = (i == 0) # The first item in a path is usually a parent menu
                            next_child = click_path[1] if is_parent and len(click_path) > 1 else None
                            
                            # Execute the click
                            smart_click(page, step, is_parent_menu=is_parent, child_text=next_child)

                        page.wait_for_timeout(3000) # Allow page content to render fully
                        
                        # 3. Screenshot Capture Sequence
                        print("       -> Capturing Desktop Screenshot...")
                        captured_images = []
                        
                        if page_info.get("needs_scroll"):
                            # Logic for long pages: Capture top, scroll down, capture bottom.
                            
                            # Top image
                            img_top = os.path.join(temp_img_dir, f"{server_name}_step_{page_index}_top.png")
                            ImageGrab.grab().save(img_top)
                            captured_images.append(img_top)
                            
                            # Focus browser center and simulate PageDown keypresses
                            page.mouse.click(browser_w // 2, screen_h // 2)
                            page.keyboard.press("PageDown")
                            page.keyboard.press("PageDown")
                            page.wait_for_timeout(1500) 
                            
                            # Bottom image
                            img_bot = os.path.join(temp_img_dir, f"{server_name}_step_{page_index}_bottom.png")
                            ImageGrab.grab().save(img_bot)
                            captured_images.append(img_bot)
                        else:
                            # Logic for short pages: Single capture
                            img_single = os.path.join(temp_img_dir, f"{server_name}_step_{page_index}.png")
                            ImageGrab.grab().save(img_single)
                            captured_images.append(img_single)
                        
                        # Store results in the main dictionary
                        captured_data[page_index][server_name] = captured_images
                        
                except Exception as e:
                    print(f"[ERROR] Failure processing {server_name}. Error: {str(e)}")
                    # Record the error string instead of an image path if something failed
                    for i in range(len(NAVIGATION_PAGES)):
                        if server_name not in captured_data[i]: captured_data[i][server_name] = f"ERROR: {str(e)}"
                finally:
                    context.close() # Cleanly close the session for this server
            browser.close() # Close browser when all servers are processed

    except Exception as general_error:
        print(f"\n[CRITICAL ERROR] Script crashed: {str(general_error)}. Generating Word Document...")
    
    finally:
        # ==============================================================================
        # POST-EXECUTION & REPORT GENERATION
        # Runs regardless of script success/failure to ensure data is preserved.
        # ==============================================================================
        print("\n[INFO] Compiling Word Document...")
        
        # Terminate the specific PowerShell window created at start
        os.system(f'taskkill /F /FI "WINDOWTITLE eq {CMD_UNIQUE_TITLE}*" >nul 2>&1')
        
        # Initialize Word Document
        doc = docx.Document()
        for section in doc.sections:
            section.left_margin, section.right_margin = Inches(0.5), Inches(0.5)

        # Add Title and Metadata Table
        doc.add_heading('Nova-HUB iDrac 10 ATP', level=1)
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.cell(0,0).text, table.cell(0,1).text = "PO", PO_INPUT
        table.cell(1,0).text, table.cell(1,1).text = "SO", SO_INPUT
        table.cell(2,0).text, table.cell(2,1).text = "SN", SN_INPUT
        
        doc.add_paragraph("")
        doc.add_heading('Chapter 1: iDRAC Verification', level=1)

        # Iterate through captured data to populate document
        for page_index, page_info in enumerate(NAVIGATION_PAGES):
            doc.add_heading(f"{page_index + 1}. {page_info['name']}:", level=2)
            
            for server_name in SERVERS.keys():
                p_server = doc.add_paragraph()
                p_server.add_run(f"{server_name}:").bold = True
                
                result = captured_data[page_index].get(server_name)
                
                # Check if data is an image path list or an error string
                if isinstance(result, list):
                    for img_path in result:
                        if os.path.exists(img_path): 
                            doc.add_picture(img_path, width=Inches(7.5)) # Resize to fit page width
                elif isinstance(result, str) and result.startswith("ERROR:"): 
                    doc.add_paragraph(result)
                else: 
                    doc.add_paragraph("No screenshot captured.")

        report_filename = f"NovaHUB_iDRAC10_ATP_{current_time}.docx"
        try:
            doc.save(report_filename)
            print(f"\n[SUCCESS] Final Report saved as: {report_filename}")
        except Exception as e: 
            print(f"\n[ERROR] Failed to save Word document. Error: {e}")

if __name__ == "__main__":
    run_automation()