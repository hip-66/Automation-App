# ==============================================================================
# Script Name: iDRAC9_Report.py
# Target: iDRAC 9 Servers ONLY
# Description: Widens browser to 72% to force standard Top Menu navigation, 
#              avoiding the hamburger menu collapse. Includes 3x zoom and 
#              double-screenshot scrolling logic for specific pages.
# ==============================================================================

import os
import sys
import subprocess
import time
import ctypes
from ctypes import wintypes
from datetime import datetime

# ==============================================================================
# USER INPUTS: Collect metadata for the final Word document.
# ==============================================================================
print("="*50)
print(" iDRAC 9 ATP Report Generator ")
print("="*50)
PO_INPUT = input("Please enter PO: ").strip()  # Purchase Order Number
SO_INPUT = input("Please enter SO: ").strip()  # Sales Order Number
SN_INPUT = input("Please enter SN: ").strip()  # System Serial Number

# ==============================================================================
# OFFLINE PACKAGE INSTALLATION
# Ensures required Python libraries are installed from a local directory 
# without needing an active internet connection.
# ==============================================================================
OFFLINE_PACKAGES_PATH = r"G:\Automation Library No Internet\offline_packages"

def install_package(pkg):
    """
    Checks if a Python module is available. If not, installs it using pip 
    from a specified offline local directory.
    """
    try: __import__(pkg if pkg != "Pillow" else "PIL")
    except: subprocess.check_call([sys.executable, "-m", "pip", "install", "--no-index", f"--find-links={OFFLINE_PACKAGES_PATH}", pkg])

# Verify and install dependencies
install_package("python-docx")    # For creating Word documents
install_package("playwright")     # For browser automation
install_package("Pillow")         # For capturing screenshots

import docx
from docx.shared import Inches, Pt
from playwright.sync_api import sync_playwright
from PIL import ImageGrab

# ==============================================================================
# CONFIGURATION VARIABLES
# Defines the target servers and the navigation flow.
# ==============================================================================
SERVERS = {
    "SRVMGT": "192.168.80.127",
    "NGINX": "192.168.80.128"
}

USERNAME = "root"
PASSWORD = "admin1234"
CMD_UNIQUE_TITLE = "" # Placeholder for dynamically generated CMD window title

# Paths adjusted for standard top-menu navigation in iDRAC 9.
# 'needs_scroll' or 'needs_double_pic': Flags that indicate whether the script 
# needs to take a top screenshot, scroll down, and take a bottom screenshot 
# to capture the full page content.
NAVIGATION_PAGES = [
    {
        "name": "view entire server’s health status", 
        "click_path": ["Dashboard"], 
        "needs_scroll": False,
        "needs_double_pic": False
    },
    {
        "name": "Common Settings", 
        "click_path": ["iDRAC Settings", "Connectivity", "Network", "Common Settings"], 
        "needs_scroll": False,
        "needs_double_pic": False
    },
    {
        "name": "IPv4 Settings", 
        "click_path": ["iDRAC Settings", "Connectivity", "Network", "IPv4 Settings"], 
        "needs_scroll": False,
        "needs_double_pic": True  # Set to True to capture top and bottom
    },
    {
        "name": "Virtual Disks", 
        "click_path": ["Storage", "Overview", "Virtual Disks"], 
        "needs_scroll": False,
        "needs_double_pic": False
    },
    {
        "name": "Firmware Inventory", 
        "click_path": ["System", "Inventory", "Firmware Inventory"], 
        "needs_scroll": False,
        "needs_double_pic": False
    },
    {
        "name": "Memory Settings", 
        "click_path": ["Configuration", "BIOS Settings", "Memory Settings"], 
        "needs_scroll": False,
        "needs_double_pic": True  # Set to True to capture top and bottom
    },
    {
        "name": "Processor Settings", 
        "click_path": ["Configuration", "BIOS Settings", "Processor Settings"], 
        "needs_scroll": True,
        "needs_double_pic": True
    },
    {
        "name": "Boot Settings", 
        "click_path": ["Configuration", "BIOS Settings", "Boot Settings"], 
        "needs_scroll": False,
        "needs_double_pic": True  # Set to True to capture top and bottom
    },
    {
        "name": "Integrated Devices", 
        "click_path": ["Configuration", "BIOS Settings", "Integrated Devices"], 
        "needs_scroll": True,
        "needs_double_pic": True
    },
    {
        "name": "System Profile Settings", 
        "click_path": ["Configuration", "BIOS Settings", "System Profile Settings"], 
        "needs_scroll": True,
        "needs_double_pic": True
    }
]

# ==============================================================================
# OS-LEVEL WINDOW MANIPULATION FUNCTIONS (WINDOWS API)
# ==============================================================================

def get_screen_dimensions():
    """
    Retrieves the primary monitor's resolution (width and height) excluding 
    the taskbar area using Windows API.
    """
    user32 = ctypes.windll.user32
    rect = wintypes.RECT()
    user32.SystemParametersInfoW(48, 0, ctypes.byref(rect), 0) # 48 is SPI_GETWORKAREA
    return rect.right - rect.left, rect.bottom - rect.top

def find_window_by_title_substring(substring):
    """
    Enumerates all open windows to find a specific window handle (HWND) 
    where the window title contains the provided substring.
    """
    hwnd_list = []
    def enum_cb(hwnd, lparam):
        length = ctypes.windll.user32.GetWindowTextLengthW(hwnd)
        if length > 0:
            buff = ctypes.create_unicode_buffer(length + 1)
            ctypes.windll.user32.GetWindowTextW(hwnd, buff, length + 1)
            if substring in buff.value: hwnd_list.append(hwnd)
        return True
    ctypes.windll.user32.EnumWindows(ctypes.WINFUNCTYPE(ctypes.c_bool, ctypes.POINTER(ctypes.c_int), ctypes.POINTER(ctypes.c_int))(enum_cb), 0)
    return hwnd_list[0] if hwnd_list else None

def setup_cmd_window(target_x, target_y, target_w, target_h):
    """
    Launches a PowerShell window, retrieves IP config data, and physically 
    manipulates the window (resizing, moving, scrolling) to sit alongside 
    the browser.
    """
    global CMD_UNIQUE_TITLE
    # Create a unique title based on current time to identify the window easily
    CMD_UNIQUE_TITLE = f"iDRAC_CMD_{int(time.time())}"
    
    # Launch PowerShell with a specific title, widened buffer, and run ipconfig
    ps_command = f'start powershell -NoExit -Command "$Host.UI.RawUI.WindowTitle = \'{CMD_UNIQUE_TITLE}\'; $psw = $Host.UI.RawUI; $sz = $psw.BufferSize; $sz.Width = 200; $psw.BufferSize = $sz; Clear-Host; ipconfig /all"'
    subprocess.Popen(ps_command, shell=True)
    time.sleep(3) # Wait for PowerShell to open and process command
    
    hwnd = find_window_by_title_substring(CMD_UNIQUE_TITLE)
    if not hwnd: return # Exit if window couldn't be found
        
    # Bring the window to the foreground and maximize/restore it
    ctypes.windll.user32.ShowWindow(hwnd, 9)
    ctypes.windll.user32.SetForegroundWindow(hwnd)
    time.sleep(1)
    
    # Virtual Key codes for OS-level keyboard injection
    VK_CONTROL, VK_OEM_MINUS = 0x11, 0xBD
    
    # 1. Inject keyboard commands to Zoom Out (Ctrl + Minus) twice
    for _ in range(2):
        ctypes.windll.user32.keybd_event(VK_CONTROL, 0, 0, 0)
        ctypes.windll.user32.keybd_event(VK_OEM_MINUS, 0, 0, 0)
        time.sleep(0.05)
        ctypes.windll.user32.keybd_event(VK_OEM_MINUS, 0, 2, 0) # Key Up
        ctypes.windll.user32.keybd_event(VK_CONTROL, 0, 2, 0)   # Key Up
        time.sleep(0.3)
        
    # 2. Relocate and resize the CMD window to the right side of the screen
    ctypes.windll.user32.SetWindowPos(hwnd, 0, int(target_x), int(target_y), int(target_w), int(target_h), 0x0040)
    time.sleep(1)
    
    # 3. Physically move the mouse and click the center of the CMD window to force focus
    click_x, click_y = int(target_x + (target_w / 2)), int(target_y + (target_h / 2))
    ctypes.windll.user32.SetCursorPos(click_x, click_y)
    ctypes.windll.user32.mouse_event(0x0002, 0, 0, 0, 0) # Left Mouse Down
    ctypes.windll.user32.mouse_event(0x0004, 0, 0, 0, 0) # Left Mouse Up
    time.sleep(0.5)
    
    # 4. Scroll the CMD window up using simulated mouse wheel events
    for _ in range(15): 
        ctypes.windll.user32.mouse_event(0x0800, 0, 0, 120 * 5, 0) # 0x0800 is MOUSEEVENTF_WHEEL
        time.sleep(0.05)

# ==============================================================================
# STANDARD NAVIGATION ENGINE
# ==============================================================================
def smart_click(page, step_text):
    """
    Intelligently locates and clicks menu items on the page.
    Attempts exact matches first, then partial matches.
    """
    print(f"          > Attempting click: '{step_text}'")
    
    # Attempt 1: Look for exact text match
    try:
        loc = page.get_by_text(step_text, exact=True).filter(visible=True)
        if loc.count() > 0:
            loc.first.click(timeout=3000)
            page.wait_for_timeout(1500) # Wait for animation/load
            return True
    except: pass
    
    # Attempt 2: Look for partial text match (contains) as a fallback
    try:
        loc = page.locator(f"//*[normalize-space(text())='{step_text}']").filter(visible=True)
        if loc.count() > 0:
            loc.first.click(timeout=3000)
            page.wait_for_timeout(1500)
            return True
    except: pass
    
    print(f"          [WARN] Failed to click: '{step_text}'")
    return False

# ==============================================================================
# MAIN AUTOMATION WORKFLOW
# ==============================================================================
def run_automation():
    current_time = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    temp_img_dir = "temp_screenshots" # Directory to temporarily store screenshots
    if not os.path.exists(temp_img_dir): os.makedirs(temp_img_dir)
    
    # Dictionary to store paths of captured images, organized by page index
    captured_data = {i: {} for i in range(len(NAVIGATION_PAGES))}
    screen_w, screen_h = get_screen_dimensions()
    
    # *** CRITICAL CONFIGURATION: Widen Browser to 72% ***
    # This specific width forces iDRAC 9 to display the standard top menu 
    # instead of collapsing into the responsive "hamburger" menu.
    browser_w = int(screen_w * 0.72)
    cmd_w = screen_w - browser_w

    try:
        # Prepare the external CMD window
        setup_cmd_window(browser_w, 0, cmd_w, screen_h)

        # Launch the browser using Playwright
        with sync_playwright() as p:
            browser = p.chromium.launch(
                headless=False, # Show the browser UI
                channel="chrome", 
                args=[f'--window-position=0,0', f'--window-size={browser_w},{screen_h}']
            )
            
            for server_name, ip_address in SERVERS.items():
                print(f"\n[INFO] Processing iDRAC 9 Server: {server_name} ({ip_address})...")
                context = browser.new_context(ignore_https_errors=True, viewport={'width': browser_w, 'height': screen_h})
                page = context.new_page()
                
                try:
                    # ==========================================================
                    # LOGIN SEQUENCE
                    # ==========================================================
                    page.goto(f"https://{ip_address}/", wait_until="domcontentloaded", timeout=60000)
                    
                    # Wait for username field to appear
                    user_sel = 'input[name="username"], input[name="user"], #username, #user'
                    page.wait_for_selector(user_sel, state="visible", timeout=20000)
                    
                    # Enter credentials and click Login
                    page.locator(user_sel).first.click() 
                    page.keyboard.type(USERNAME, delay=50) 
                    page.locator('input[name="password"], #password').first.click()
                    page.keyboard.type(PASSWORD, delay=50)
                    page.locator('button:has-text("Log In"), #btn-login, input[type="submit"]').first.click()
                    
                    # ==========================================================
                    # POST-LOGIN SEQUENCE
                    # ==========================================================
                    print("       -> Waiting 5 seconds after login...")
                    page.wait_for_timeout(5000) # Mandatory wait to ensure full page load

                    print("       -> Applying 3x Zoom Out (Ctrl + -)...")
                    # Click top-left safely to focus browser without clicking UI elements
                    page.mouse.click(10, 10) 
                    # Inject 3 Zoom Out commands to ensure all content fits on screen
                    for _ in range(3):
                        page.keyboard.down('Control')
                        page.keyboard.press('-')
                        page.keyboard.up('Control')
                        page.wait_for_timeout(500)

                    # ==========================================================
                    # NAVIGATION SEQUENCE
                    # ==========================================================
                    for page_index, page_info in enumerate(NAVIGATION_PAGES):
                        print(f"\n       -> Step {page_index + 1}: {page_info['name']}")
                        
                        # Reset view to Dashboard before starting a new path (if not the first step)
                        if page_index > 0:
                            try:
                                page.locator('text="Dashboard"').first.click(timeout=3000)
                                page.wait_for_timeout(1500)
                            except: pass
                        
                        # Iterate through the defined steps for the current page
                        click_path = page_info['click_path']
                        for step in click_path:
                            if step == "Dashboard" and page_index == 0: continue
                            smart_click(page, step)

                        page.wait_for_timeout(3000) # Allow page content to render fully
                        
                        # ==========================================================
                        # SCREENSHOT CAPTURE SEQUENCE
                        # ==========================================================
                        print("       -> Capturing Screenshot(s)...")
                        imgs = []
                        
                        # Check if this specific page requires capturing both top and bottom views
                        if page_info.get("needs_double_pic") or page_info.get("needs_scroll"):
                            # 1. Top Screenshot
                            img_top = os.path.join(temp_img_dir, f"{server_name}_{page_index}_top.png")
                            ImageGrab.grab().save(img_top)
                            imgs.append(img_top)
                            
                            # 2. Scroll down using JavaScript (avoids clicking the screen and losing focus)
                            page.evaluate("window.scrollBy(0, 1000)")
                            page.wait_for_timeout(1500) # Wait for smooth scroll
                            
                            # 3. Bottom Screenshot
                            img_bot = os.path.join(temp_img_dir, f"{server_name}_{page_index}_bot.png")
                            ImageGrab.grab().save(img_bot)
                            imgs.append(img_bot)
                        else:
                            # Single Screenshot
                            img_s = os.path.join(temp_img_dir, f"{server_name}_{page_index}.png")
                            ImageGrab.grab().save(img_s)
                            imgs.append(img_s)
                        
                        # Store results in the main dictionary
                        captured_data[page_index][server_name] = imgs
                        
                except Exception as e:
                    print(f"       [ERROR] Failed on {server_name}: {e}")
                    # Record the error string instead of an image path if something failed
                    for i in range(len(NAVIGATION_PAGES)):
                        if server_name not in captured_data[i]: captured_data[i][server_name] = f"ERROR: {e}"
                finally:
                    context.close() # Cleanly close the session for this server
            browser.close() # Close browser when all servers are processed
            
    except Exception as e:
        print(f"\n[CRITICAL] Automation interrupted: {e}")
    
    finally:
        # ==============================================================================
        # POST-EXECUTION & REPORT GENERATION
        # Runs regardless of script success/failure to ensure data is preserved.
        # ==============================================================================
        print("\n[INFO] Compiling Word Document...")
        
        # Terminate the specific PowerShell window created at start
        os.system(f'taskkill /F /FI "WINDOWTITLE eq {CMD_UNIQUE_TITLE}*" >nul 2>&1')
        
        # Initialize Word Document
        doc = docx.Document()
        for section in doc.sections:
            section.left_margin, section.right_margin = Inches(0.5), Inches(0.5)

        # Add Title and Metadata Table
        doc.add_heading('Nova-HUB iDrac 9 ATP Report', level=1)
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.cell(0,0).text, table.cell(0,1).text = "PO", PO_INPUT
        table.cell(1,0).text, table.cell(1,1).text = "SO", SO_INPUT
        table.cell(2,0).text, table.cell(2,1).text = "SN", SN_INPUT
        doc.add_paragraph("")

        # Iterate through captured data to populate document
        for page_index, page_info in enumerate(NAVIGATION_PAGES):
            doc.add_heading(f"{page_index + 1}. {page_info['name']}:", level=2)
            
            for server_name in SERVERS.keys():
                p = doc.add_paragraph()
                p.add_run(f"{server_name}:").bold = True
                
                res = captured_data[page_index].get(server_name)
                
                # Check if data is an image path list or an error string
                if isinstance(res, list):
                    for img in res:
                        if os.path.exists(img): 
                            doc.add_picture(img, width=Inches(7.5)) # Resize to fit page width
                elif isinstance(res, str): 
                    doc.add_paragraph(res)
                else: 
                    doc.add_paragraph("No data captured.")

        report_name = f"NovaHUB_iDRAC9_ATP_{current_time}.docx"
        try: 
            doc.save(report_name)
            print(f"[SUCCESS] Report saved: {report_name}")
        except Exception as e: 
            print(f"[ERROR] Save failed: {e}")

if __name__ == "__main__":
    run_automation()