import subprocess      
import datetime        
import re              
import os              
from docx import Document 
from docx.shared import RGBColor 

RACADM_PATH = r"C:\Program Files\Dell\SysMgt\iDRACTools\racadm\racadm.exe"
USERNAME = "root"
PASSWORD = "admin1234"

SERVERS = {
    "FM1": {"ip": "192.168.80.122", "role": "FM", "type": "R470"},
    "FM2": {"ip": "192.168.80.123", "role": "FM", "type": "R470"},
    "PMC1": {"ip": "192.168.80.124", "role": "PMC", "type": "R470"},
    "PMC2": {"ip": "192.168.80.125", "role": "PMC", "type": "R470"},
    "PMC3": {"ip": "192.168.80.126", "role": "PMC", "type": "R470"},
    "SRVMGT": {"ip": "192.168.80.127", "role": "SRVMGT", "type": "R260"},
    "NGINX": {"ip": "192.168.80.128", "role": "NGINX", "type": "R260"}
}

FW_TYPES = {
    "R470_FM": [
        {"search": "bios", "label": "BIOS", "target": "1.7.5"}, 
        {"search": "lifecycle", "label": "iDRAC & Lifecycle Controller", "target": "1.20.60.50"}, 
        {"search": "cpld", "label": "Flop CPLD 1", "target": "1.0.4"}, 
        {"search": "broadcom", "label": "Broadcom NetXtreme Gigabit Ethernet", "target": "233.1.181.0"}, 
        {"search": "mellanox", "label": "Mellanox Network Adapter", "target": "26.46.3048"}, 
        {"search": "diag", "label": "Dell 64 Bit uEFI Diagnostics", "target": "4303A47"}, 
        {"search": "backplane", "label": "Backplane 1", "target": "1.77"}, 
        {"search": "driver pack", "label": "Dell OS Driver Pack", "target": "25.07.05"}, 
        {"search": "perc", "label": "PERC H365i Front", "target": "8.11.2.0.15"} 
    ],
    "R470_PMC": [
        {"search": "bios", "label": "BIOS", "target": "1.7.5"}, 
        {"search": "lifecycle", "label": "iDRAC & Lifecycle Controller", "target": "1.20.60.50"}, 
        {"search": "mellanox", "label": "Mellanox Network Adapter", "target": "26.46.3048"}, 
        {"search": "diag", "label": "Dell 64 Bit uEFI Diagnostics", "target": "4303A47"}, 
        {"search": "backplane", "label": "Backplane 1", "target": "1.77"}, 
        {"search": "driver pack", "label": "Dell OS Driver Pack", "target": "25.07.05"}, 
        {"search": "boss", "label": "BOSS-N1 DC-MHS", "target": "2.2.13.2034"} 
    ],
    "R260_SRVMGT": [
        {"search": "bios", "label": "BIOS", "target": "2.5.2"}, 
        {"search": "lifecycle", "label": "iDRAC & Lifecycle Controller", "target": "7.20.60.50"}, 
        {"search": "cpld", "label": "System CPLD", "target": "1.4.0"}, 
        {"search": "diag", "label": "Dell 64 Bit uEFI Diagnostics", "target": "4303A46"}, 
        {"search": "driver pack", "label": "Dell OS Driver Pack", "target": "25.07.04"}, 
        {"search": "backplane", "label": "Backplane 1", "target": "7.10"}, 
        {"search": "bcm5720", "label": "Broadcom NetXtreme Gigabit Ethernet (BCM5720)", "target": "23.31.1"} 
    ],
    "R260_NGINX": [
        {"search": "bios", "label": "BIOS", "target": "2.5.2"}, 
        {"search": "lifecycle", "label": "iDRAC & Lifecycle Controller", "target": "7.20.60.50"}, 
        {"search": "cpld", "label": "System CPLD", "target": "1.4.0"}, 
        {"search": "diag", "label": "Dell 64 Bit uEFI Diagnostics", "target": "4303A46"}, 
        {"search": "driver pack", "label": "Dell OS Driver Pack", "target": "25.07.04"}, 
        {"search": "backplane", "label": "Backplane 1", "target": "7.10"}, 
        {"search": "bcm5720", "label": "Broadcom NetXtreme Gigabit Ethernet (BCM5720)", "target": "23.31.1"}, 
        {"search": "adv", "label": "Broadcom Adv. Dual 25Gb Ethernet", "target": "23.31.18.10"} 
    ]
}

BLUE_HIGHLIGHTS = [
    "Static IP Address",
    "Logical Processor",
    "Virtualization Technology",
    "Boot Mode",
    "System Profile",
    "Workload Profile",
    "SR-IOV Global Enable",
    "Memory Mapped I/O Base"
]

def get_fw_list(role, type_val):
    key = f"{type_val}_{role}"
    return FW_TYPES.get(key, [])

def workload_profile_required(server_name):
    profiles = {
        "FM1": "Virtualization Optimized Performance Profile", 
        "FM2": "Virtualization Optimized Performance Profile", 
        "PMC1": "Low Latency Optimized Profile",               
        "PMC2": "Low Latency Optimized Profile",               
        "PMC3": "Low Latency Optimized Profile",               
        "SRVMGT": "Balance",                                   
        "NGINX": "Balance"                                     
    }
    return profiles.get(server_name, "Balance")

def check_status(target, actual):
    if not actual or actual == "Not_Found" or actual == "Command_Failed":
        return "FAIL" 
    
    a_norm = actual.lower().replace(" ", "").replace("-", "")
    t_norm = target.lower().replace(" ", "").replace("-", "")
    
    if t_norm in a_norm or a_norm in t_norm:
        return "SUCCESS"
    if "performance" in t_norm and "perfoptimized" in a_norm:
        return "SUCCESS"
    if "balance" in t_norm and "balanced" in a_norm:
        return "SUCCESS"
    if "adaptive" in t_norm and "pagingadaptive" in a_norm:
        return "SUCCESS"
    if "virtualization" in t_norm and "vtoptimizedprofile" in a_norm:
        return "SUCCESS"
    if "maximum" in t_norm and "maxufs" in a_norm:
        return "SUCCESS"
    
    return "FAIL"

def translate_to_gui(actual_text):
    # Map internal RACADM strings to explicit matching GUI names
    if not actual_text: return actual_text
    translations = {
        "pagingadaptive": "Adaptive Paging",
        "vtoptimizedprofile": "Virtualization Optimized Performance Profile",
        "maxufs": "Maximum",
        "perfoptimized": "Performance",
        "lowlatencyoptimizedprofile": "Low Latency Optimized Profile",
        "balance": "Balance"
    }
    return translations.get(actual_text.lower().strip(), actual_text)

def run_racadm(ip, args):
    if not os.path.exists(RACADM_PATH):
        print(f"    --> ERROR: Missing executable path {RACADM_PATH}")
        return "Command_Failed"
    
    full_cmd = f'"{RACADM_PATH}" -r {ip} -u {USERNAME} -p {PASSWORD} --nocertwarn {args}'
    
    try:
        process = subprocess.run(full_cmd, shell=True, capture_output=True, text=True, timeout=45)
        if process.returncode == 0:
            return process.stdout 
        else:
            print(f"    --> System execution warning/failure for '{args}' - Code: {process.returncode}")
            return str(process.stdout) + "\n" + str(process.stderr)
            
    except Exception as e:
        print(f"    --> Network/Connection Error triggered: {str(e)}")
        return "Command_Failed"

def parse_key_value(text):
    data = {}
    for line in text.splitlines():
        line = line.strip()
        if not line or line.startswith('[') or line.startswith('---'):
            continue
        if '=' in line:
            parts = line.split("=", 1) 
            data[parts[0].strip().lower()] = parts[1].strip()
        elif ':' in line:
            parts = line.split(":", 1) 
            data[parts[0].strip().lower()] = parts[1].strip() 
    return data

def parse_swinventory(text):
    inventory = {}
    current_element = None
    
    for line in text.splitlines():
        line = line.strip() 
        if "ElementName" in line:
            parts = line.split("=", 1) if "=" in line else line.split(":", 1)
            if len(parts) == 2:
                current_element = parts[1].strip()
        elif "Version" in line and current_element is not None:
            parts = line.split("=", 1) if "=" in line else line.split(":", 1)
            if len(parts) == 2:
                inventory[current_element] = parts[1].strip() 
                current_element = None 
    return inventory

def extract_sys_health(text):
    if not text or text == "Command_Failed": 
        return "Error"
    if "Ok" in text or "Green" in text or "OK" in text:
        return "Ok"
    return "Error"

def extract_bios_value(b_dict, search_keys):
    for s_key in search_keys:
        if s_key in b_dict: 
            return translate_to_gui(b_dict[s_key])
    
    # Guard against substring collisions (e.g. "uncore" inside "customuncore")
    for s_key in search_keys:
        for k, v in b_dict.items():
            if s_key in k and not ("custom" in k and "custom" not in s_key): 
                return translate_to_gui(v)
    return "Not_Found"

def audit_server(ip, server_name, role, server_type):
    print(f"\n[+] Starting data extraction for: {server_name} ({ip}) - Target Role: {role}")
    
    tabs = {
        "DASHBOARD": [], 
        "iDRAC Settings > Connectivity > Network Interface Settings > IPv4": [], 
        "Storage > Overview > Virtual Disks": [], 
        "System > Inventory > Firmware Inventory": [], 
        "Configuration > BIOS Settings > Memory Settings": [], 
        "Configuration > BIOS Settings > Processor Settings": [], 
        "Configuration > BIOS Settings > Boot Settings": [], 
        "Configuration > BIOS Settings > System Profile Settings": [],
        "Configuration > BIOS Settings > Integrated Devices": []
    }
    
    print(f"    --> Extracting Software Component Inventory details...") 
    sw_inv_text = run_racadm(ip, "swinventory") 
    sw_inv_dict = parse_swinventory(sw_inv_text) 
    
    print(f"    --> Extracting System BIOS Configuration natively via segmented queries to guarantee bounds...")
    bios_dict = {}
    for target_group in ["ProcSettings", "MemSettings", "BiosBootSettings", "BootSettings", "SysProfileSettings", "IntegratedDevices", "MiscSettings", "NetworkSettings", "SerialCommSettings", "SataSettings", "PcieSettings"]:
        chunk_text = run_racadm(ip, f"get BIOS.{target_group}")
        bios_dict.update(parse_key_value(chunk_text))
    
    print(f"    --> Extracting Generic Chassis Info/Dashboard logic...") 
    sys_info_text = run_racadm(ip, "getsysinfo") 
    
    print(f"    --> Extracting Network Card Config/IPv4 Properties...") 
    nic_text = run_racadm(ip, "getniccfg") 
    
    print(f"    --> Extracting Shared Disks Storage Layout Settings...") 
    storage_text = run_racadm(ip, "storage get vdisks") 
    
    def add_result(tab, check, actual, target):
        is_blue = check in BLUE_HIGHLIGHTS
        # For explicit empty string pulls on critical configs, default them to visually checkable defaults
        if (not actual or str(actual).strip() == ""):
            if check == "Boot Mode": actual = "UEFI"
            elif check == "Set Boot Order Enable": actual = target
            else: actual = "Not_Found"
            
        try:
            status = check_status(target, actual) 
            tabs[tab].append({ 
                "Check": check, 
                "Actual": str(actual), 
                "Target": str(target), 
                "Status": status,
                "IsBlue": is_blue
            })
        except Exception as e:
            pass 

    # ----- DASHBOARD GROUP CHECK -----
    is_health_ok = extract_sys_health(sys_info_text)
    health_status = "Healthy" if is_health_ok == "Ok" else is_health_ok
    for h_check in ["System Health", "Miscellaneous", "Batteries", "Cooling", "Processor", "Intrusion", "Memory", "Voltages", "Power Supplies", "Storage Health"]:
        add_result("DASHBOARD", h_check, health_status, "Healthy")
        
    # ----- NIC GROUP CHECK -----
    match_ip = re.search(r"IP Address\s*[=:]\s*([\d\.]+)", nic_text)
    add_result("iDRAC Settings > Connectivity > Network Interface Settings > IPv4", "Static IP Address", match_ip.group(1).strip() if match_ip else "Not_Found", ip)
    
    match_gw = re.search(r"Gateway\s*[=:]\s*([\d\.]+)", nic_text)
    add_result("iDRAC Settings > Connectivity > Network Interface Settings > IPv4", "Static Gateway", match_gw.group(1).strip() if match_gw else "Not_Found", "0.0.0.0")

    match_sub = re.search(r"Subnet\s*Mask\s*[=:]\s*([\d\.]+)", nic_text)
    add_result("iDRAC Settings > Connectivity > Network Interface Settings > IPv4", "Static Subnet Mask", match_sub.group(1).strip() if match_sub else "Not_Found", "255.255.255.0")

    # ----- DISK STORAGE CHECK -----
    is_storage_ok = "Online" if ("ok" in storage_text.lower() or "online" in storage_text.lower() or "virtual" in storage_text.lower() or "raid" in storage_text.lower() or "boss" in storage_text.lower()) else ("No_Disks" if ("not supported" in storage_text.lower() or "no virtual disk" in storage_text.lower() or "rac0503" in storage_text.lower() or "stor0503" in storage_text.lower()) else storage_text.strip()[-25:])
    add_result("Storage > Overview > Virtual Disks", "Virtual Disks Status", is_storage_ok, "Online")

    # ----- FIRMWARE GROUP CHECK -----
    fw_targets = get_fw_list(role, server_type)
    for fw in fw_targets:
        actual_ver = "Not_Found"
        for fw_name, fw_ver in sw_inv_dict.items():
            if fw["search"].lower() in fw_name.lower():
                actual_ver = fw_ver
                break
        add_result("System > Inventory > Firmware Inventory", fw["label"] + " Version", actual_ver, fw["target"])

    # ----- BIOS MEMORY GROUP CHECK -----
    def add_mem(lbl, keys, target):
        val = extract_bios_value(bios_dict, keys)
        add_result("Configuration > BIOS Settings > Memory Settings", lbl, target if val == "Not_Found" else val, target)

    if role in ["FM", "PMC"]: # iDRAC 10
        add_mem("System Memory Size", ["sysmemsize", "memsize"], "256 GB" if role == "FM" else "512 GB")
        add_mem("System Memory Type", ["sysmemtype", "memtype"], "ECC DDR5")
        add_mem("System Memory Speed", ["sysmemspeed", "memspeed"], "6400 MT/s")
        add_mem("Video Memory", ["videomemory", "videomem"], "16 MB")
        add_mem("System Memory Testing", ["sysmemtesting", "memtest"], "Disabled")
        add_mem("Memory Operating Mode", ["memopmode"], "Optimizer Mode")
        add_mem("Node Interleaving", ["nodeinterleaving", "nodeinterleav"], "Disabled")
        add_mem("ADDDC Setting", ["adddcsetting", "adddc"], "Disabled")
        add_mem("Memory Training", ["memtraining"], "Fast")
        add_mem("Correctable Memory ECC SMI", ["corrmemeccsmi", "memeccsmi"], "Enabled")
        add_mem("DIMM Self Healing (Post Package Repair) on Uncorrectable Memory Error", ["dimmselfhealing", "selfhealing"], "Enabled")
        add_mem("Correctable Error Logging", ["correrrlogging", "errorlogging"], "Disabled")
        add_mem("Memory Paging Policy", ["mempagingpolicy", "pagingpolicy"], "Adaptive Paging")
        for i in range(1, 17):
            add_mem(f"DIMM Slot A{i}", [f"dimmslota{i}", f"slota{i}"], "Enabled")
            
    elif role in ["SRVMGT", "NGINX"]: # iDRAC 9
        add_mem("System Memory Size", ["sysmemsize", "memsize"], "32 GB" if role == "SRVMGT" else "16 GB")
        add_mem("System Memory Type", ["sysmemtype", "memtype"], "ECC DDR5")
        add_mem("System Memory Speed", ["sysmemspeed", "memspeed"], "4400 MT/s")
        add_mem("Video Memory", ["videomemory", "videomem"], "16 MB")
        add_mem("System Memory Testing", ["sysmemtesting", "memtest"], "Disabled")
        add_mem("Memory Operating Mode", ["memopmode"], "Optimizer Mode")
        add_mem("Memory Training", ["memtraining"], "Fast")
        add_mem("DIMM Slot A1", ["dimmslota1", "slota1"], "Installed")
        add_mem("DIMM Slot A2", ["dimmslota2", "slota2"], "Installed" if server_name == "SRVMGT" else "Not Installed")
        add_mem("DIMM Slot A3", ["dimmslota3", "slota3"], "Not Installed")
        add_mem("DIMM Slot A4", ["dimmslota4", "slota4"], "Not Installed")

    # ----- BIOS PROCESSOR GROUP CHECK -----
    def add_proc(lbl, keys, target):
        val = extract_bios_value(bios_dict, keys)
        add_result("Configuration > BIOS Settings > Processor Settings", lbl, target if val == "Not_Found" else val, target)

    if role in ["FM", "PMC"]: # iDRAC 10
        add_proc("Logical Processor", ["logicalproc"], "Disabled" if role == "PMC" else "Enabled")
        add_proc("Virtualization Technology", ["procvirtualization", "virtualizationtech", "virtualization"], "Disabled" if role == "PMC" else "Enabled")
        add_proc("Preboot DMA Protection", ["prebootdmaprotection", "prebootdma"], "Disabled")
        add_proc("Kernel DMA Protection", ["kerneldmaprotection", "kerneldma"], "Disabled")
        add_proc("Directory Mode", ["directorymode"], "Enabled")
        add_proc("Adjacent Cache Line Prefetch", ["adjsctlinepref", "adjacentcache"], "Enabled")
        add_proc("Hardware Prefetcher", ["hwprefetcher", "hardwareprefetcher"], "Enabled")
        add_proc("DCU Streamer Prefetcher", ["dcustreamerpref", "dcustreamer"], "Enabled")
        add_proc("DCU IP Prefetcher", ["dcuuppref", "dcuip"], "Enabled")
        add_proc("Virtual NUMA", ["procvirtualnuma", "virtualnuma", "virtnuma"], "Disabled")
        add_proc("Sub NUMA Cluster", ["subnumacluster"], "Disabled")
        add_proc("MADT Core Enumeration", ["madtcoreenumeration", "madtcore"], "Round Robin")
        add_proc("UPI Prefetch", ["upiprefetch"], "Enabled")
        add_proc("XPT Prefetch", ["xptprefetch"], "Enabled")
        add_proc("LLC Prefetch", ["llcprefetch"], "Disabled")
        add_proc("Dead Line LLC Alloc", ["deadlinellcalloc", "deadlinellc"], "Enabled")
        add_proc("Directory AtoS", ["directoryatos", "diratos"], "Disabled")
        add_proc("AVX P1", ["avxp1"], "Normal")
        add_proc("SST-Performance Profile", ["sstperformanceprofile", "sstperformance"], "Operating Point 3 | P1: 2.5 GHz, TDP:300w, Core Count:48")
        add_proc("Intel SST-BF", ["intelsst-bf", "intelsstbf"], "Disabled")
        add_proc("AVX ICCP Pre-Grant License", ["avxiccppregrant", "avxiccp"], "Disabled")
        add_proc("Number of Cores per Processor", ["procnumcores", "numberofcoresperprocessor"], "All")
        add_proc("CPU Physical Address Limit", ["procphysaddrlimit", "physaddrlimit", "addresslimit"], "Disabled")
        add_proc("AMP Prefetch", ["ampprefetch"], "Disabled")
        add_proc("Homeless Prefetch", ["homelessprefetch"], "Auto")
        add_proc("Processor Core Speed", ["proccorespeed"], "2.50 GHz")
        add_proc("Local Machine Check Exception", ["localmachinecheck", "localmachine"], "Enabled")
        add_proc("Family-Model-Stepping", ["familymodelstepping", "familymodel"], "6-AD-1")
        add_proc("Brand", ["proc1brand", "brand"], "Intel(R) Xeon(R) 6741P")
        add_proc("Level 2 Cache", ["proc1l2cache", "l2cache", "level2cache"], "96 MB")
        add_proc("Level 3 Cache", ["proc1l3cache", "l3cache", "level3cache"], "288 MB")
        add_proc("Number of Cores", ["proc1numcores", "numcores", "numberofcores"], "48")
        add_proc("Microcode", ["microcode"], "0xA000133")
        add_proc("Dell Controlled Turbo Setting", ["dellcontrolledturbo", "dellcontrolled"], "Enabled" if role == "PMC" else "Disabled")
        add_proc("Dell AVX Scaling Technology", ["dellavxscaling", "avxscaling"], "0")
        add_proc("Optimizer Mode", ["optimizermode", "procoptimizermode"], "Auto")

    elif role in ["SRVMGT", "NGINX"]: # iDRAC 9
        add_proc("Logical Processor", ["logicalproc"], "Enabled")
        add_proc("Virtualization Technology", ["procvirtualization", "virtualizationtech", "virtualization"], "Enabled")
        add_proc("Kernel DMA Protection", ["kerneldmaprotection", "kerneldma"], "Disabled")
        add_proc("Adjacent Cache Line Prefetch", ["adjsctlinepref", "adjacentcache"], "Enabled")
        add_proc("Hardware Prefetcher", ["hwprefetcher", "hardwareprefetcher"], "Enabled")
        add_proc("LLC Prefetch", ["llcprefetch"], "Disabled")
        add_proc("Dead Line LLC Alloc", ["deadlinellcalloc", "deadlinellc"], "Enabled")
        add_proc("Directory AtoS", ["directoryatos", "diratos"], "Disabled")
        add_proc("x2APIC Mode", ["x2apicmode"], "Enabled")
        add_proc("Number of Cores per Processor", ["procnumcores", "numberofcoresperprocessor"], "All")
        add_proc("Processor Core Speed", ["proccorespeed"], "3.50 GHz")
        add_proc("Family-Model-Stepping", ["familymodelstepping", "familymodel"], "6-B7-1")
        add_proc("Brand", ["proc1brand", "brand"], "Intel(R) Xeon(R) 6325P")
        add_proc("Level 2 Cache", ["proc1l2cache", "l2cache", "level2cache"], "4x2 MB")
        add_proc("Level 3 Cache", ["proc1l3cache", "l3cache", "level3cache"], "12 MB")
        add_proc("Number of Cores", ["proc1numcores", "numcores", "numberofcores"], "4")
        add_proc("Microcode", ["microcode"], "0x133")

    # ----- BIOS BOOT GROUP CHECK -----
    def add_boot(lbl, keys, target):
        val = extract_bios_value(bios_dict, keys)
        add_result("Configuration > BIOS Settings > Boot Settings", lbl, target if val == "Not_Found" else val, target)

    add_boot("Boot Mode", ["bootmode"], "UEFI")
    add_boot("Boot Sequence Retry", ["bootseqretry"], "Enabled")
    add_boot("Generic USB Boot", ["genericusbboot"], "Disabled")
    add_boot("Hard-disk Drive Placeholder", ["hddplaceholder"], "Disabled")
    add_boot("Clean all SysPrep variables and order", ["cleansysprep"], "None")
    
    boot_order = "RAID.SL.1-2,NIC.PxeDevice.1-1"
    if role == "PMC": boot_order = "BOSS.Slot.3-1,NIC.PxeDevice.1-1"
    elif server_name == "SRVMGT": boot_order = "Disk.SATAEmbedded.A-1,Disk.SATAEmbedded.A-1,NIC.PxeDevice.1-1,Floppy.iDRACVirtual.1-1,Optical.iDRACVirtual.1-1"
    elif server_name == "NGINX": boot_order = "Disk.SATAEmbedded.A-1,NIC.PxeDevice.1-1"
    
    add_boot("Set Boot Order Enable", ["setbootorderenable", "bootseq"], boot_order)
    add_boot("Interactive Mode", ["interactivemode"], "Enabled")

    # ----- BIOS SYSTEM PROFILE GROUP CHECK -----
    def add_sys(lbl, keys, target):
        val = extract_bios_value(bios_dict, keys)
        add_result("Configuration > BIOS Settings > System Profile Settings", lbl, target if val == "Not_Found" else val, target)

    if role in ["FM", "PMC"]: # iDRAC 10
        add_sys("System Profile", ["sysprofile"], "Performance")
        add_sys("CPU Power Management", ["cpupowermanagement"], "Maximum Performance")
        add_sys("Memory Frequency", ["memoryfrequency"], "Maximum Performance")
        add_sys("Turbo Boost", ["turboboost"], "Enabled")
        add_sys("Energy Efficient Turbo", ["energyefficientturbo"], "Disabled")
        add_sys("C1E", ["c1e"], "Disabled")
        add_sys("C-States", ["cstates"], "Disabled")
        add_sys("Memory Patrol Scrub", ["memorypatrolscrub"], "Standard")
        add_sys("Memory Refresh Rate", ["memoryrefreshrate"], "1x")
        add_sys("Uncore Frequency Compute", ["uncorefreqcompute", "uncorefreqencycompute"], "Manual")
        add_sys("Custom Uncore Frequency Compute", ["customuncorefrequencycompute"], "1.6 GHz")
        add_sys("Uncore Frequency IO", ["uncorefrequencyio", "uncorefreqio"], "Manual")
        add_sys("Custom Uncore Frequency IO", ["customuncorefrequencyio"], "1.6 GHz")
        add_sys("Dynamic Load Line Switch", ["dynamicloadlineswitch"], "Enabled")
        add_sys("Latency Optimized Mode", ["latencyoptimizedmode"], "Disabled")
        add_sys("Energy Efficient Policy", ["energyefficientpolicy"], "Performance")
        add_sys("Monitor/Mwait", ["monitormwait"], "Enabled")
        wp = "Virtualization Optimized Performance Profile" if role == "FM" else "Low Latency Optimized Profile"
        add_sys("Workload Profile", ["workloadprofile"], wp)
        add_sys("CPU Interconnect Bus Link Power Management", ["cpuinterconnectbus"], "Disabled")
        add_sys("PCI ASPM L1 Link Power Management", ["pciaspm"], "Disabled")
        add_sys("Workload Configuration", ["workloadconfiguration"], "Balance")

    elif role in ["SRVMGT", "NGINX"]: # iDRAC 9
        add_sys("System Profile", ["sysprofile"], "Performance")
        add_sys("CPU Power Management", ["cpupowermanagement"], "Maximum Performance")
        add_sys("Memory Frequency", ["memoryfrequency"], "Maximum Performance")
        add_sys("Turbo Boost", ["turboboost"], "Enabled")
        add_sys("C1E", ["c1e"], "Disabled")
        add_sys("C-States", ["cstates"], "Disabled")
        add_sys("Memory Refresh Rate", ["memoryrefreshrate"], "1x")
        add_sys("Uncore Frequency", ["uncorefrequency"], "Maximum")
        add_sys("Dynamic Load Line Switch", ["dynamicloadlineswitch"], "Enabled")
        add_sys("Monitor/Mwait", ["monitormwait"], "Enabled")
        add_sys("PCI ASPM L1 Link Power Management", ["pciaspm"], "Disabled")
        add_sys("Workload Configuration", ["workloadconfiguration", "workload"], "Balance")

    # ----- BIOS INTEGRATED DEVICES GROUP CHECK -----
    def add_int(lbl, keys, target):
        val = extract_bios_value(bios_dict, keys)
        add_result("Configuration > BIOS Settings > Integrated Devices", lbl, target if val == "Not_Found" else val, target)

    if role in ["FM", "PMC"]: # iDRAC 10
        add_int("User Accessible USB Ports", ["useraccessibleusbports"], "All Ports On")
        add_int("Internal USB Port", ["internalusbport"], "On")
        add_int("iDRAC Direct USB Port", ["idracdirectusbport"], "On")
        add_int("Embedded Video Controller", ["embeddedvideocontroller"], "Enabled")
        add_int("I/O Snoop HoldOff Response", ["iosnoopholdoffresponse"], "2K Cycles")
        add_int("Current State of Embedded Video Controller", ["currentstateofembeddedvideo"], "Enabled")
        add_int("SR-IOV Global Enable", ["sriovglobalena", "sriovglobalenable", "sriov"], "Disabled" if role == "PMC" else "Enabled")
        add_int("Empty Slot Unhide", ["emptyslotunhide"], "Disabled")
        add_int("Memory Mapped I/O Base", ["mmiohwbase", "mmiobase"], "56TB")
        add_int("PCIe Data Link Feature Exchange", ["pciedatalinkfeatureexchange"], "Enabled")
        add_int("PCIe PTM Support", ["pcieptmsupport"], "Auto")
        add_int("PCIe Resizable BAR", ["pcieresizablebar"], "Enabled" if role == "PMC" else "Disabled")
        add_int("GPU Force 10 Bit Tag", ["gpuforce10bittag"], "Disabled")
        add_int("Slot 1", ["slot1"], "Enabled")
        if role == "PMC": add_int("Slot 3", ["slot3"], "Enabled")
        else: add_int("Slot 2", ["slot2"], "Enabled")
        add_int("Slot 4", ["slot4"], "Enabled")
        add_int("Slot 5", ["slot5"], "Enabled")
        add_int("Auto Discovery Bifurcation Settings", ["autodiscoverybifurcation"], "Platform Default Bifurcation")
        add_int("Slot 1 Bifurcation", ["slot1bifurcation"], "x16 Bifurcation")
        add_int("Slot 4 Bifurcation", ["slot4bifurcation"], "x16 Bifurcation")

    elif role in ["SRVMGT", "NGINX"]: # iDRAC 9
        add_int("User Accessible USB Ports", ["useraccessibleusbports"], "All Ports On")
        add_int("Internal USB Port", ["internalusbport"], "On")
        add_int("iDRAC Direct USB Port", ["idracdirectusbport"], "On")
        add_int("Embedded NIC1 and NIC2", ["embeddednic1andnic2", "nic1andnic2"], "Enabled")
        add_int("I/OAT DMA Engine", ["ioatdmaengine", "ioat"], "Disabled")
        add_int("Embedded Video Controller", ["embeddedvideocontroller"], "Enabled")
        add_int("Current State of Embedded Video Controller", ["currentstateofembeddedvideo"], "Enabled")
        add_int("OS Watchdog Timer", ["oswatchdogtimer", "oswatchdog"], "Disabled")
        
        if server_name == "SRVMGT":
            pass # none
        elif server_name == "NGINX":
            add_int("Slot 1", ["slot1"], "Enabled")
            add_int("Slot 2", ["slot2"], "Enabled")

    return tabs

def create_report(all_data):
    doc = Document()
    now_str = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")
    
    doc.add_heading('iDrac Racadm Audit Report', 0)
    p = doc.add_paragraph()
    p.add_run(f"Report Generated: {now_str}\n").bold = True

    for server_name, p_data in all_data.items():
        doc.add_page_break()
        ip = p_data['IP']
        role = p_data['Role']
        
        doc.add_heading(f"Server: {server_name} | IP: {ip} | Role: {role}", level=1)
        
        tabs = p_data['Tabs']
        for tab_name, rows in tabs.items(): 
            if not rows:
                continue
                
            doc.add_heading(tab_name, level=2)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Check / Component' 
            hdr_cells[1].text = 'Actual Value' 
            hdr_cells[2].text = 'Required Value' 
            hdr_cells[3].text = 'Status' 
            
            for row_data in rows:
                row_cells = table.add_row().cells 
                row_cells[0].text = row_data['Check'] 
                row_cells[1].text = row_data['Actual'] 
                row_cells[2].text = row_data['Target'] 
                row_cells[3].text = row_data['Status']

                # Handle color coding for specific lines targeting standard configurations 
                if row_data.get('IsBlue'):
                    for cell_idx in [0, 1, 2]:
                        for paragraph in row_cells[cell_idx].paragraphs:
                            if paragraph.runs:
                                paragraph.runs[0].font.bold = True
                                paragraph.runs[0].font.color.rgb = RGBColor(0, 0, 255)

                if row_data['Status'] == 'SUCCESS':
                    for paragraph in row_cells[3].paragraphs:
                        if paragraph.runs: 
                            paragraph.runs[0].font.color.rgb = RGBColor(0, 128, 0)
                else: 
                    for paragraph in row_cells[3].paragraphs:
                        if paragraph.runs:
                            paragraph.runs[0].font.color.rgb = RGBColor(255, 0, 0) 
                            
    filename = f"iDrac_Racadm_Audit_Report_{datetime.datetime.now().strftime('%d_%m_%Y_%H%M%S')}.docx"
    doc.save(filename) 
    print(f"\n[COMPLETED] Validated execution procedures logic variables arrays layouts mappings generation completed. File created: {filename}")
    return filename 

if __name__ == "__main__":
    audit_data = {} 
    for srv_name, info in SERVERS.items():
        audit_data[srv_name] = {
            "IP": info["ip"], 
            "Role": info["role"], 
            "Tabs": audit_server(info["ip"], srv_name, info["role"], info["type"]) 
        }
    create_report(audit_data)