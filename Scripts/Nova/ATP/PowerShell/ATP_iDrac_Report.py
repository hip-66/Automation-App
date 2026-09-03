# ==============================================================================
# Script Name: NovaHUB_Combined_iDRAC_Report.py
# Target: iDRAC 9 & iDRAC 10 Servers
# Description: Unified automation script to connect to iDRAC servers, navigate 
#              through configuration pages, capture screenshots, and compile a 
#              SINGLE Word document report. 
#              NOTE: iDRAC 9 browser width adjusted to match iDRAC 10 (65%).
# Environment: Offline
# ==============================================================================

import os
import sys
import subprocess
import time
import ctypes
import shutil
from ctypes import wintypes
from datetime import datetime

# ==============================================================================
# USER INPUTS: Collected once at the start of the process
# ==============================================================================
print("="*50)
print(" Nova-HUB Unified iDRAC ATP Report Generator ")
print("="*50)
# metadata for the Word document table
PO_INPUT = input("Please enter PO: ").strip()  # Purchase Order Number
SO_INPUT = input("Please enter SO: ").strip()  # Sales Order Number
SN_INPUT = input("Please enter SN: ").strip()  # System Serial Number
print("="*50)

# ==============================================================================
# OFFLINE PACKAGE INSTALLATION: Ensures required libraries are present
# ==============================================================================
OFFLINE_PACKAGES_PATH = r"G:\Automation Library No Internet\offline_packages"

def install_package_offline(package_name, import_name):
    """ Checks for a module and installs it from local storage if missing. """
    try:
        __import__(import_name)
    except ImportError:
        print(f"[INFO] Installing '{package_name}'...")
        try:
            # Install using local wheels path without internet access
            subprocess.check_call([sys.executable, "-m", "pip", "install", "--no-index", f"--find-links={OFFLINE_PACKAGES_PATH}", package_name])
        except subprocess.CalledProcessError:
            print(f"[ERROR] Failed to install '{package_name}'.")

# Dependency Verification
install_package_offline("python-docx", "docx")    # For Word report creation
install_package_offline("playwright", "playwright") # For browser automation
install_package_offline("Pillow", "PIL")          # For desktop screenshots

import docx
from docx.shared import Inches, Pt
from playwright.sync_api import sync_playwright
from PIL import ImageGrab

# ==============================================================================
# GLOBAL CONFIGURATION & SERVER DEFINITIONS
# ==============================================================================
USERNAME = "root"
PASSWORD = "admin1234"
TEMP_IMG_DIR = "temp_screenshots" # Directory for temporary image storage

# Ensure the temp directory exists
if not os.path.exists(TEMP_IMG_DIR):
    os.makedirs(TEMP_IMG_DIR)

# Targeted Servers - Grouped by Generation
SERVERS_10 = {
    "FM1": "192.168.80.122",
    "FM2": "192.168.80.123",
    "PMC1": "192.168.80.124",
    "PMC2": "192.168.80.125",
    "PMC3": "192.168.80.126"
}

SERVERS_9 = {
    "SRVMGT": "192.168.80.127",
    "NGINX": "192.168.80.128"
}

# Navigation Paths: Preserved exactly from original source files
NAV_PAGES_10 = [
    {"name": "iDRAC view entire server's health status", "click_path": ["Dashboard"], "needs_scroll": False},
    {"name": "Server Names", "click_path": ["iDRAC Settings", "Connectivity", "Network", "Common Settings"], "needs_scroll": False},
    {"name": "IPv4 Settings", "click_path": ["iDRAC Settings", "Connectivity", "Network Interface Settings", "IPv4"], "needs_scroll": False},
    {"name": "Virtual Disks", "click_path": ["Storage", "Overview", "Virtual Disks"], "needs_scroll": False},
    {"name": "Firmware Inventory", "click_path": ["System", "Inventory", "Firmware Inventory"], "needs_scroll": False},
    {"name": "Memory setting", "click_path": ["Configuration", "BIOS Settings", "Memory Settings"], "needs_scroll": False},
    {"name": "Processor setting", "click_path": ["Configuration", "BIOS Settings", "Processor Settings"], "needs_scroll": True},
    {"name": "Boot setting", "click_path": ["Configuration", "BIOS Settings", "Boot Settings"], "needs_scroll": False},
    {"name": "FM Integrated Devices", "click_path": ["Configuration", "BIOS Settings", "More", "Integrated Devices"], "needs_scroll": True},
    {"name": "System profiles settings", "click_path": ["Configuration", "BIOS Settings", "More", "System Profile Settings"], "needs_scroll": True}
]

NAV_PAGES_9 = [
    {"name": "view entire server’s health status", "click_path": ["Dashboard"], "needs_double_pic": False},
    {"name": "Common Settings", "click_path": ["iDRAC Settings", "Connectivity", "Network", "Common Settings"], "needs_double_pic": False},
    {"name": "IPv4 Settings", "click_path": ["iDRAC Settings", "Connectivity", "Network", "IPv4 Settings"], "needs_double_pic": True},
    {"name": "Virtual Disks", "click_path": ["Storage", "Overview", "Virtual Disks"], "needs_double_pic": False},
    {"name": "Firmware Inventory", "click_path": ["System", "Inventory", "Firmware Inventory"], "needs_double_pic": False},
    {"name": "Memory Settings", "click_path": ["Configuration", "BIOS Settings", "Memory Settings"], "needs_double_pic": True},
    {"name": "Processor Settings", "click_path": ["Configuration", "BIOS Settings", "Processor Settings"], "needs_double_pic": True},
    {"name": "Boot Settings", "click_path": ["Configuration", "BIOS Settings", "Boot Settings"], "needs_double_pic": True},
    {"name": "Integrated Devices", "click_path": ["Configuration", "BIOS Settings", "Integrated Devices"], "needs_double_pic": True},
    {"name": "System Profile Settings", "click_path": ["Configuration", "BIOS Settings", "System Profile Settings"], "needs_double_pic": True}
]

# ==============================================================================
# OS-LEVEL WINDOW MANIPULATION FUNCTIONS (WINDOWS API)
# ==============================================================================

def get_screen_dimensions():
    """ Returns monitor width/height excluding taskbar using Windows SystemParametersInfo. """
    user32 = ctypes.windll.user32
    rect = wintypes.RECT()
    user32.SystemParametersInfoW(48, 0, ctypes.byref(rect), 0) # SPI_GETWORKAREA
    return rect.right - rect.left, rect.bottom - rect.top

def find_window_by_title_substring(substring):
    """ Scans all active windows to locate a specific handle (HWND) by name. """
    hwnd_list = []
    def enum_cb(hwnd, lparam):
        length = ctypes.windll.user32.GetWindowTextLengthW(hwnd)
        if length > 0:
            buff = ctypes.create_unicode_buffer(length + 1)
            ctypes.windll.user32.GetWindowTextW(hwnd, buff, length + 1)
            if substring in buff.value: hwnd_list.append(hwnd)
        return True
    ctypes.windll.user32.EnumWindows(ctypes.WINFUNCTYPE(ctypes.c_bool, ctypes.POINTER(ctypes.c_int), ctypes.POINTER(ctypes.c_int))(enum_cb), 0)
    return hwnd_list[0] if hwnd_list else None

def setup_cmd_window(target_x, target_y, target_w, target_h):
    """ Opens PowerShell, runs ipconfig, and positions the window beside the browser. """
    unique_title = f"iDRAC_CMD_{int(time.time())}"
    # Start PowerShell with a specific title and ipconfig /all
    ps_command = f'start powershell -NoExit -Command "$Host.UI.RawUI.WindowTitle = \'{unique_title}\'; $psw = $Host.UI.RawUI; $sz = $psw.BufferSize; $sz.Width = 200; $psw.BufferSize = $sz; Clear-Host; ipconfig /all"'
    subprocess.Popen(ps_command, shell=True)
    time.sleep(3)
    
    hwnd = find_window_by_title_substring(unique_title)
    if hwnd:
        ctypes.windll.user32.ShowWindow(hwnd, 9) # Restore window
        ctypes.windll.user32.SetForegroundWindow(hwnd) # Bring to front
        time.sleep(1)
        # Apply Ctrl+Minus zoom out via keyboard events
        VK_CONTROL, VK_OEM_MINUS = 0x11, 0xBD
        for _ in range(2):
            ctypes.windll.user32.keybd_event(VK_CONTROL, 0, 0, 0)
            ctypes.windll.user32.keybd_event(VK_OEM_MINUS, 0, 0, 0)
            time.sleep(0.05)
            ctypes.windll.user32.keybd_event(VK_OEM_MINUS, 0, 2, 0)
            ctypes.windll.user32.keybd_event(VK_CONTROL, 0, 2, 0)
            time.sleep(0.3)
        # Position window on the right side of the screen
        ctypes.windll.user32.SetWindowPos(hwnd, 0, int(target_x), int(target_y), int(target_w), int(target_h), 0x0040)
        time.sleep(1)
        # Move mouse and click to lock focus on CMD
        click_x, click_y = int(target_x + (target_w / 2)), int(target_y + (target_h / 2))
        ctypes.windll.user32.SetCursorPos(click_x, click_y)
        ctypes.windll.user32.mouse_event(0x0002, 0, 0, 0, 0) # Down
        ctypes.windll.user32.mouse_event(0x0004, 0, 0, 0, 0) # Up
        time.sleep(0.5)
        # Physical scroll to top via mouse wheel events
        for _ in range(15):
            ctypes.windll.user32.mouse_event(0x0800, 0, 0, 120 * 5, 0)
            time.sleep(0.05)
    return unique_title

# ==============================================================================
# SMART CLICK ENGINES: Preserves original logic for menu navigation
# ==============================================================================

def smart_click_10(page, step_text, is_parent=False, child=None):
    """ Handles iDRAC 10 specific menu interactions and expansions. """
    if is_parent and child:
        try:
            if page.locator(f'text="{child}"').first.is_visible(timeout=1000): return True 
        except: pass
    try:
        loc = page.get_by_text(step_text, exact=True).filter(visible=True)
        if loc.count() > 0:
            loc.first.click(timeout=3000); page.wait_for_timeout(1500); return True
    except: pass
    try:
        loc = page.locator(f"//*[contains(text(), '{step_text}')]").filter(visible=True)
        if loc.count() > 0:
            loc.first.click(timeout=3000); page.wait_for_timeout(1500); return True
    except: pass
    return False

def smart_click_9(page, step_text):
    """ Handles iDRAC 9 specific menu interactions for top-menu layout. """
    try:
        loc = page.get_by_text(step_text, exact=True).filter(visible=True)
        if loc.count() > 0:
            loc.first.click(timeout=3000); page.wait_for_timeout(1500); return True
    except: pass
    try:
        loc = page.locator(f"//*[normalize-space(text())='{step_text}']").filter(visible=True)
        if loc.count() > 0:
            loc.first.click(timeout=3000); page.wait_for_timeout(1500); return True
    except: pass
    return False

# ==============================================================================
# MAIN WORKFLOW EXECUTION
# ==============================================================================

def run_automation():
    """ Orchestrates the connection, navigation, and document building. """
    timestamp = datetime.now().strftime("%d-%m-%Y_%H-%M")
    report_filename = f"NovaHUB_ATP_Report_{timestamp}.docx"
    screen_w, screen_h = get_screen_dimensions()
    
    # Unified results data structure: results[test_index][server_name] = [images]
    captured_results = {i: {} for i in range(10)}

    # Global Width for all iDRAC generations (Adjusted to 65% for uniformity)
    global_browser_w = int(screen_w * 0.65)
    global_cmd_w = screen_w - global_browser_w

    # CMD Initialization - Run ONCE for the entire process
    cmd_title = setup_cmd_window(global_browser_w, 0, global_cmd_w, screen_h)

    try:
        with sync_playwright() as p:
            # --- PHASE 1: iDRAC 10 SERVERS ---
            print("\n>>> Starting iDRAC 10 servers batch...")
            browser10 = p.chromium.launch(headless=False, channel="chrome", args=[f'--window-position=0,0', f'--window-size={global_browser_w},{screen_h}'])
            
            for s_name, ip in SERVERS_10.items():
                print(f"[INFO] Connecting to iDRAC 10: {s_name} ({ip})")
                context = browser10.new_context(ignore_https_errors=True, viewport={'width': global_browser_w, 'height': screen_h})
                page = context.new_page()
                try:
                    page.goto(f"https://{ip}/", timeout=60000)
                    # Login handling for iDRAC 10
                    u_sel = 'input[name="username"], input[name="user"], #username, #user, #idrac_user'
                    page.wait_for_selector(u_sel, state="visible", timeout=20000)
                    page.locator(u_sel).first.fill(USERNAME)
                    page.locator('input[name="password"], #password, #idrac_password').first.fill(PASSWORD)
                    page.locator('button:has-text("Log In"), button:has-text("Login"), #btn-login').first.click()
                    page.wait_for_timeout(5000)
                    
                    # Manual Zoom Out for UI visibility
                    page.mouse.click(global_browser_w // 2, screen_h // 2)
                    page.keyboard.down('Control'); page.keyboard.press('-'); page.keyboard.up('Control')
                    
                    for idx, nav in enumerate(NAV_PAGES_10):
                        try: page.locator('text="Dashboard"').first.click(timeout=2000)
                        except: pass
                        for i, step in enumerate(nav['click_path']):
                            if step == "Dashboard": continue
                            smart_click_10(page, step, (i==0), nav['click_path'][1] if len(nav['click_path'])>1 else None)
                        page.wait_for_timeout(3000)
                        
                        img = os.path.join(TEMP_IMG_DIR, f"10_{s_name}_{idx}.png")
                        ImageGrab.grab().save(img); captured_results[idx][s_name] = [img]
                        
                        if nav['needs_scroll']:
                            page.keyboard.press("PageDown"); page.wait_for_timeout(1500)
                            img_sc = os.path.join(TEMP_IMG_DIR, f"10_{s_name}_{idx}_sc.png")
                            ImageGrab.grab().save(img_sc); captured_results[idx][s_name].append(img_sc)
                except Exception as e: print(f"[ERROR] {s_name}: {e}")
                finally: context.close()
            browser10.close()

            # --- PHASE 2: iDRAC 9 SERVERS ---
            print("\n>>> Starting iDRAC 9 servers batch...")
            # Browser size for iDRAC 9 now matches iDRAC 10 (global_browser_w)
            browser9 = p.chromium.launch(headless=False, channel="chrome", args=[f'--window-position=0,0', f'--window-size={global_browser_w},{screen_h}'])
            
            for s_name, ip in SERVERS_9.items():
                print(f"[INFO] Connecting to iDRAC 9: {s_name} ({ip})")
                context = browser9.new_context(ignore_https_errors=True, viewport={'width': global_browser_w, 'height': screen_h})
                page = context.new_page()
                try:
                    # ORIGINAL LOGIN LOGIC FOR iDRAC 9 (Corrected based on provided screenshot)
                    page.goto(f"https://{ip}/", wait_until="domcontentloaded", timeout=60000)
                    u_sel_9 = 'input[name="username"], input[name="user"], #username, #user'
                    page.wait_for_selector(u_sel_9, state="visible", timeout=20000)
                    page.locator(u_sel_9).first.click()
                    page.keyboard.type(USERNAME, delay=50)
                    page.locator('input[name="password"], #password').first.click()
                    page.keyboard.type(PASSWORD, delay=50)
                    page.locator('button:has-text("Log In"), #btn-login, input[type="submit"]').first.click()
                    
                    page.wait_for_timeout(6000)
                    # Manual Zoom Out for iDRAC 9 legacy UI
                    page.mouse.click(10, 10)
                    for _ in range(3): page.keyboard.down('Control'); page.keyboard.press('-'); page.keyboard.up('Control'); time.sleep(0.5)
                    
                    for idx, nav in enumerate(NAV_PAGES_9):
                        if idx > 0:
                            try: page.locator('text="Dashboard"').first.click(timeout=3000)
                            except: pass
                        for step in nav['click_path']:
                            if step == "Dashboard" and idx == 0: continue
                            smart_click_9(page, step)
                        page.wait_for_timeout(3000)
                        
                        img = os.path.join(TEMP_IMG_DIR, f"9_{s_name}_{idx}.png")
                        ImageGrab.grab().save(img); captured_results[idx][s_name] = [img]
                        
                        if nav['needs_double_pic']:
                            page.evaluate("window.scrollBy(0, 1000)"); page.wait_for_timeout(1500)
                            img_sc = os.path.join(TEMP_IMG_DIR, f"9_{s_name}_{idx}_sc.png")
                            ImageGrab.grab().save(img_sc); captured_results[idx][s_name].append(img_sc)
                except Exception as e: print(f"[ERROR] {s_name}: {e}")
                finally: context.close()
            browser9.close()

    finally:
        # --- DOCUMENT COMPILATION ---
        print("\n>>> Building Unified Word Document...")
        doc = docx.Document()
        for section in doc.sections: section.left_margin = section.right_margin = Inches(0.5)
        
        doc.add_heading('Nova-HUB iDRAC ATP Report', level=1)
        tbl = doc.add_table(rows=3, cols=2); tbl.style = 'Table Grid'
        tbl.cell(0,0).text, tbl.cell(0,1).text = "PO", PO_INPUT
        tbl.cell(1,0).text, tbl.cell(1,1).text = "SO", SO_INPUT
        tbl.cell(2,0).text, tbl.cell(2,1).text = "SN", SN_INPUT
        doc.add_paragraph("")

        # Global Order: FM1-FM2 -> PMC1-PMC3 -> SRVMGT -> NGINX
        all_ordered_servers = list(SERVERS_10.keys()) + list(SERVERS_9.keys())
        
        for idx in range(10):
            test_title = NAV_PAGES_10[idx]['name']
            doc.add_heading(f"{idx + 1}. {test_title}:", level=2)
            for s_name in all_ordered_servers:
                if s_name in captured_results[idx]:
                    doc.add_paragraph(f"{s_name}:").runs[0].bold = True
                    for img_path in captured_results[idx][s_name]:
                        if os.path.exists(img_path): doc.add_picture(img_path, width=Inches(7.5))

        # Save and Cleanup
        try:
            doc.save(report_filename)
            print(f"\n[SUCCESS] Final Report saved: {report_filename}")
        except:
            print(f"\n[ERROR] Failed to save document.")
            
        # Physical Task Kill of the initialized CMD session
        os.system(f'taskkill /F /FI "WINDOWTITLE eq {cmd_title}*" >nul 2>&1')
        # Recursive removal of temp image directory
        if os.path.exists(TEMP_IMG_DIR): shutil.rmtree(TEMP_IMG_DIR)

if __name__ == "__main__":
    run_automation()